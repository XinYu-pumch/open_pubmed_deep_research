
import os
import asyncio
import logging
import re
import json
import time
import random
import sys
import requests
import xml.etree.ElementTree as ET
import httpx
import numpy as np
import chromadb
import shutil
import io
import glob
import pandas as pd
import zipfile
import locale
from urllib.parse import urlencode
from bs4 import BeautifulSoup
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Import workspace manager
from workspace_manager import (
    get_app_resource_path,
    get_collection_path,
    get_content_dir,
    get_same_directory_converter_executable,
    is_bundled_app,
)

try:
    from curl_cffi.requests import AsyncSession
    HAS_CURL_CFFI = True
except ImportError:
    HAS_CURL_CFFI = False

logger = logging.getLogger(__name__)

# ==========================================
#               1. Constants & Prompts
# ==========================================
SCIHUB_MIRRORS = [
    "https://sci-hub.st",
    "https://sci-hub.red",
    "https://sci-hub.box",
    "https://sci-hub.ru",
    "https://sci-hub.se",
    "https://sci-hub.ee",
    "https://sci-hub.wf",
    "https://sci-hub.yt",
]

# Sample DOIs for testing mirror availability
TEST_DOIS = [
    "10.1007/s102270000022",
    "10.1097/01.tp.0000186382.81130.ba",
    "10.2217/bmm-2020-0073",
]

# Available mirrors cache (mirrors that passed testing)
_AVAILABLE_MIRRORS_CACHE: List[str] = []
_MIRRORS_TESTED = False  # Flag indicating whether mirrors have been tested

CONTENT_CHUNK_SIZE = 8192 
MIN_CHUNK_CHARS = 100
ALL_ABSTRACTS_DOC_ID = "doc_all_abstracts"
COMBINED_ABSTRACT_METADATA_KEY = "pmid_abstract_pairs_json"
TITLE_PMID_DOI_DOC_ID = "title_pmid_doi"
TITLE_PMID_DOI_METADATA_KEY = "title_pmid_doi_json"

FRAMEWORK_PROMPT_INTRO_ZH = "现在我想根据一些我检索的文献撰写综述，我已经汇总给你文献的摘要合集。\n我希望你根据这些摘要合集写一个综述框架，主题为{topic}。\n"
FRAMEWORK_PROMPT_INTRO_EN = "I want to write a literature review based on the papers I retrieved. I have provided a collection of abstracts.\nPlease generate a review framework for the topic: {topic}.\n"

FRAMEWORK_PROMPT_BASE_REQS_ZH = """要求：
- 用中文撰写，尽可能分点，涵盖的维度尽可能详实；
- 综述框架的结构为多级标题，第一级标题为阿拉伯数字1、3、4这样，第二级为1.1、1.2、2.3这样，第三级为1.1.1、2.1.3这样；
- 第一级标题通常较为简洁，从第二级开始，标题需要是较为完整的句子；
- 各级标题彼此独立，互不遗漏；"""

FRAMEWORK_PROMPT_BASE_REQS_EN = """Requirements:
- Write in English, be as detailed as possible, covering comprehensive dimensions;
- Use a multi-level heading structure (e.g., 1, 1.1, 1.1.1);
- Level 1 headings should be concise; from Level 2 onwards, headings should be complete sentences;
- Headings should be mutually exclusive and collectively exhaustive;"""

FRAMEWORK_PROMPT_OUTPUT_FMT_ZH = """
最后输出的格式为csv形式，第一列为序号，第二列为标题名称。你的输出结果应该只包括这个表格，不要有表格以外的其他内容。
csv格式示例：
1,引言
1.1,研究背景
...
附摘要内容：
{abstracts}
"""

FRAMEWORK_PROMPT_OUTPUT_FMT_EN = """
Output format must be CSV, with the first column as 'Serial' and the second as 'Title'. Output ONLY the table.
CSV example:
1,Introduction
1.1,Research Background
...
Abstracts:
{abstracts}
"""

FRAMEWORK_PROMPT_INTRO = FRAMEWORK_PROMPT_INTRO_ZH
FRAMEWORK_PROMPT_BASE_REQS = FRAMEWORK_PROMPT_BASE_REQS_ZH
FRAMEWORK_PROMPT_OUTPUT_FMT = FRAMEWORK_PROMPT_OUTPUT_FMT_ZH

def get_framework_filename(lang: str = 'zh') -> str:
    return "review_framework_eng.csv" if lang == 'en' else "review_framework_cn.csv"

def get_framework_csv_path(collection_name: str, lang: str = 'zh', allow_legacy: bool = False) -> str:
    c_dir = get_collection_path(collection_name)
    preferred = os.path.join(c_dir, get_framework_filename(lang))
    if allow_legacy and not os.path.exists(preferred):
        legacy = os.path.join(c_dir, "reveiw_framework.csv")
        if os.path.exists(legacy):
            return legacy
    return preferred

def get_review_parts_dir(collection_name: str, lang: str = 'zh', allow_legacy: bool = False) -> str:
    c_dir = get_collection_path(collection_name)
    preferred_name = "review_parts_eng" if lang == 'en' else "review_parts_cn"
    preferred = os.path.join(c_dir, preferred_name)
    if allow_legacy and not os.path.exists(preferred):
        legacy = os.path.join(c_dir, "review_parts")
        if os.path.exists(legacy):
            return legacy
    return preferred

# ==========================================
#               2. Basic Helpers
# ==========================================

def make_request_with_retry(url: str, params: Dict[str, Any], method: str = 'get', max_retries: int = 3):
    for attempt in range(max_retries):
        try:
            if method.lower() == 'get': response = requests.get(url, params=params, timeout=30.0)
            else: response = requests.post(url, data=params, timeout=30.0)
            response.raise_for_status()
            return response
        except Exception:
            if attempt == max_retries - 1: raise
            time.sleep(1)

async def download_file(url: str, local_path: str, client: httpx.AsyncClient) -> bool:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
        async with client.stream("GET", url, headers=headers, timeout=60.0, follow_redirects=True) as response:
            if response.status_code != 200: return False
            with open(local_path, "wb") as f:
                async for chunk in response.aiter_bytes(): f.write(chunk)
        if os.path.exists(local_path) and os.path.getsize(local_path) < 1000:
            os.remove(local_path); return False
        return True
    except: return False

def parse_pubmed_xml_details(xml_content: bytes) -> List[Dict[str, Any]]:
    results = []
    if not xml_content: return results
    try:
        root = ET.fromstring(xml_content)
        for article in root.findall('.//PubmedArticle'):
            entry = {'pmid': None, 'title': 'N/A', 'abstract': None, 'doi': None, 'pmcid': None}
            pmid_elem = article.find('./MedlineCitation/PMID')
            if pmid_elem is not None: entry['pmid'] = pmid_elem.text.strip()
            title_elem = article.find('./MedlineCitation/Article/ArticleTitle')
            if title_elem is not None: entry['title'] = "".join(title_elem.itertext()).strip()
            abstract_node = article.find('./MedlineCitation/Article/Abstract')
            if abstract_node is not None:
                texts = ["".join(t for t in abst.itertext()).strip() for abst in abstract_node.findall('./AbstractText')]
                if texts: entry['abstract'] = "\n".join(texts)
            for item in article.findall('./PubmedData/ArticleIdList/ArticleId'):
                if item.get('IdType') == 'doi': entry['doi'] = item.text.strip()
                elif item.get('IdType') == 'pmc': entry['pmcid'] = item.text.strip()
            if entry['pmid']: results.append(entry)
    except Exception as e: logger.error(f"XML Parse Error: {e}")
    return results

def get_next_xml_filename(save_dir: str) -> str:
    files = glob.glob(os.path.join(save_dir, "pubmed_results_*.xml"))
    indices = [int(os.path.basename(f).replace("pubmed_results_", "").replace(".xml", "")) for f in files if os.path.basename(f).replace("pubmed_results_", "").replace(".xml", "").isdigit()]
    return f"pubmed_results_{max(indices) + 1 if indices else 1}.xml"

def merge_pubmed_xmls(save_dir: str, log_callback) -> Tuple[bool, str]:
    main_file = os.path.join(save_dir, "pubmed_results.xml")
    files = sorted(glob.glob(os.path.join(save_dir, "pubmed_results_*.xml")))
    if os.path.exists(main_file): files.insert(0, main_file)
    if not files: return False, "No XML files found."

    unique_articles = {}
    total_found = 0
    log_callback(f"Merging {len(files)} XML files...")

    try:
        for fp in files:
            try:
                for article in ET.parse(fp).getroot().findall('.//PubmedArticle'):
                    total_found += 1
                    pmid = article.find('.//PMID').text.strip()
                    if pmid not in unique_articles: unique_articles[pmid] = article
            except: pass
        
        new_root = ET.Element("PubmedArticleSet")
        for art in unique_articles.values(): new_root.append(art)
        tree = ET.ElementTree(new_root)
        if hasattr(ET, 'indent'): ET.indent(tree, space="  ", level=0)
        tree.write(main_file, encoding="utf-8", xml_declaration=True)
        
        del_count = 0
        for f in files:
            if os.path.abspath(f) != os.path.abspath(main_file):
                os.remove(f); del_count += 1
        return True, f"Merged {total_found} articles into {len(unique_articles)} unique records. Deleted {del_count} fragments."
    except Exception as e: return False, f"Merge failed: {e}"

def clear_xml_files(save_dir: str) -> str:
    count = 0
    for f in glob.glob(os.path.join(save_dir, "*.xml")):
        try: os.remove(f); count += 1
        except: pass
    return f"Cleared {count} XML files."

def extract_abstracts_to_txt(save_dir: str, log_callback) -> Tuple[bool, str]:
    xml_file = os.path.join(save_dir, "pubmed_results.xml")
    output_file = os.path.join(save_dir, "abstract_combined.txt")
    if not os.path.exists(xml_file): return False, "XML not found."
    try:
        with open(xml_file, "rb") as f: xml_content = f.read()
        articles = parse_pubmed_xml_details(xml_content)
        if not articles: return False, "No articles found."
        content_lines = []
        for p in articles:
            if p.get('abstract'):
                content_lines.append(f"PMID: {p.get('pmid')}\n{p.get('abstract')}\n{'-'*50}\n")
        with open(output_file, "w", encoding="utf-8") as f: f.write("\n".join(content_lines))
        return True, f"Extracted {len(content_lines)} abstracts."
    except Exception as e: return False, f"Error: {e}"

async def test_scihub_mirrors(log_callback) -> List[str]:
    """并行测试所有 Sci-Hub 镜像，返回可用镜像列表

    测试策略：
    - 按 DOI 轮次测试：DOI1 测试所有镜像 → DOI2 测试所有镜像 → DOI3 测试所有镜像
    - 每轮内所有镜像并行测试
    - 镜像只要在任意一轮成功，即加入可用列表
    """
    global _AVAILABLE_MIRRORS_CACHE

    if _AVAILABLE_MIRRORS_CACHE:
        log_callback(f"Using cached available mirrors: {len(_AVAILABLE_MIRRORS_CACHE)}")
        return _AVAILABLE_MIRRORS_CACHE

    if not HAS_CURL_CFFI:
        log_callback("curl_cffi is not installed; mirror testing will be skipped.")
        return SCIHUB_MIRRORS

    log_callback(f"Testing {len(SCIHUB_MIRRORS)} Sci-Hub mirrors...")
    log_callback(f"Test DOIs: {TEST_DOIS}")

    async def test_mirror_with_doi(mirror: str, doi: str) -> Tuple[str, bool, str]:
        """测试单个镜像 + 单个 DOI，返回 (镜像, 是否成功, 日志消息)"""
        mirror_name = mirror.split('.')[-1]

        try:
            async with AsyncSession(impersonate="chrome120", verify=False) as browser:
                page_url = f"{mirror}/{doi}"
                start_time = time.time()

                try:
                    resp = await browser.get(page_url, timeout=10, allow_redirects=True)
                except Exception as e:
                    return mirror, False, f"  [{mirror_name}] 连接失败 - {str(e)[:25]}"

                elapsed = time.time() - start_time

                if resp.status_code != 200:
                    return mirror, False, f"  [{mirror_name}] HTTP {resp.status_code} ({elapsed:.1f}s)"

                content_type = resp.headers.get('Content-Type', '').lower()
                if 'application/pdf' in content_type and len(resp.content) > 1000:
                    return mirror, True, f"  [{mirror_name}] PDF returned directly ({elapsed:.1f}s)"

                soup = BeautifulSoup(resp.text, 'lxml')
                pdf_url = None

                pdf_div = soup.find('div', class_='pdf')
                if pdf_div:
                    obj_tag = pdf_div.find('object', attrs={'type': 'application/pdf'})
                    if obj_tag and obj_tag.get('data'):
                        pdf_url = obj_tag['data']

                if not pdf_url:
                    download_div = soup.find('div', class_='download')
                    if download_div:
                        link = download_div.find('a', href=True)
                        if link:
                            pdf_url = link['href']

                if not pdf_url:
                    embed = soup.find('embed', attrs={'type': 'application/pdf'})
                    if embed and embed.get('src'):
                        pdf_url = embed['src']

                if not pdf_url:
                    iframe = soup.find('iframe', id='pdf')
                    if iframe and iframe.get('src'):
                        pdf_url = iframe['src']

                if not pdf_url:
                    if "captcha" in resp.text.lower():
                        return mirror, False, f"  [{mirror_name}] CAPTCHA triggered ({elapsed:.1f}s)"
                    else:
                        return mirror, False, f"  [{mirror_name}] PDF link not found ({elapsed:.1f}s)"

                # 规范化 PDF URL
                if '#' in pdf_url:
                    pdf_url = pdf_url.split('#')[0]
                if pdf_url.startswith('//'):
                    pdf_url = 'https:' + pdf_url
                elif pdf_url.startswith('/'):
                    pdf_url = mirror + pdf_url
                elif not pdf_url.startswith('http'):
                    pdf_url = mirror + '/' + pdf_url

                # 验证 PDF 可下载
                try:
                    pdf_start = time.time()
                    pdf_resp = await browser.get(pdf_url, timeout=30, allow_redirects=True)
                    pdf_elapsed = time.time() - pdf_start

                    if pdf_resp.status_code == 200 and len(pdf_resp.content) > 1000:
                        kb_size = len(pdf_resp.content) / 1024
                        return mirror, True, f"  [{mirror_name}] PDF verified ({kb_size:.0f}KB, {pdf_elapsed:.1f}s)"
                    else:
                        return mirror, False, f"  [{mirror_name}] PDF download failed (status={pdf_resp.status_code}, size={len(pdf_resp.content)})"
                except Exception as e:
                    return mirror, False, f"  [{mirror_name}] PDF download timed out - {str(e)[:20]}"

        except Exception as e:
            return mirror, False, f"  [{mirror_name}] Error - {str(e)[:25]}"

    # 记录已确认可用的镜像（一次成功即可）
    confirmed_mirrors = set()

    # 按 DOI 轮次测试
    for doi_index, doi in enumerate(TEST_DOIS):
        log_callback(f"--- Mirror test round {doi_index + 1}/3: {doi} ---")

        # 该轮所有镜像并行测试
        tasks = [test_mirror_with_doi(m, doi) for m in SCIHUB_MIRRORS]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 处理结果
        for result in results:
            if isinstance(result, Exception):
                log_callback(f"  Test error: {str(result)[:50]}")
            elif isinstance(result, tuple):
                mirror, success, log_msg = result
                log_callback(log_msg)
                if success:
                    confirmed_mirrors.add(mirror)

        # 显示当前轮次统计
        log_callback(f"--- Round {doi_index + 1} finished; available mirrors so far: {len(confirmed_mirrors)} ---")

    # 转换为列表，保持原始顺序
    available_mirrors = [m for m in SCIHUB_MIRRORS if m in confirmed_mirrors]

    _AVAILABLE_MIRRORS_CACHE = available_mirrors

    if not available_mirrors:
        log_callback("No Sci-Hub mirrors were verified; all mirrors will still be tried.")
        return SCIHUB_MIRRORS
    else:
        log_callback(f"Final available mirrors: {len(available_mirrors)} {[m.split('.')[-1] for m in available_mirrors]}")
        return available_mirrors


def create_rotated_mirror_order(mirrors: List[str], worker_index: int) -> List[str]:
    """为每个 worker 生成轮转后的镜像顺序"""
    if not mirrors:
        return mirrors
    n = len(mirrors)
    offset = worker_index % n
    return mirrors[offset:] + mirrors[:offset]


def normalize_pmcid(pmcid: Optional[str]) -> Optional[str]:
    """Normalize a PMCID to the canonical PMC12345 form."""
    clean_pmcid = str(pmcid or "").strip()
    if not clean_pmcid:
        return None
    if clean_pmcid.lower().startswith("pmc"):
        return f"PMC{clean_pmcid[3:]}"
    return f"PMC{clean_pmcid}"


def stringify_repo_version(version: Any) -> str:
    """Convert the repository version to the compact string expected by EuropePMC."""
    if isinstance(version, float) and version.is_integer():
        return str(int(version))
    if isinstance(version, int):
        return str(version)
    if isinstance(version, str):
        return version.strip() or "1"
    return "1"


def score_europepmc_repo_pdf_file(item: Dict[str, Any]) -> int:
    """Prefer likely article PDFs over proof, supplementary, or review artifacts."""
    filename = str(item.get("filename") or item.get("fileName") or "").strip().lower()
    stem = filename[:-4] if filename.endswith(".pdf") else filename
    tokens = {token for token in re.split(r"[^a-z0-9]+", stem) if token}

    score = 0
    if filename == "main.pdf":
        score += 100
    if "main" in tokens:
        score += 40
    if "article" in tokens or "manuscript" in tokens:
        score += 15
    if "nihms" in tokens or stem.startswith("nihms"):
        score += 8

    penalties = {
        "supp": 60,
        "supplement": 80,
        "supplementary": 80,
        "appendix": 60,
        "table": 40,
        "tables": 40,
        "figure": 40,
        "figures": 40,
        "proof": 50,
        "prf": 50,
        "coif": 45,
        "rc": 35,
        "author": 20,
        "authors": 20,
        "response": 30,
        "reply": 30,
        "rebuttal": 30,
        "checklist": 30,
        "reporting": 25,
        "cover": 20,
    }
    for token, penalty in penalties.items():
        if token in tokens:
            score -= penalty

    if item.get("url"):
        score += 3
    return score


def extract_ranked_europepmc_repo_pdf_files(fulltext_repo_metadata: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Return PDF file metadata sorted so the main article is tried before supplements/proofs."""
    if not fulltext_repo_metadata:
        return []

    files = fulltext_repo_metadata.get("files") or []
    if isinstance(files, dict):
        files = [files]

    ranked_files = []
    for index, item in enumerate(files):
        if not isinstance(item, dict):
            continue
        filename = str(item.get("filename") or item.get("fileName") or "").strip()
        mime_type = str(item.get("mimeType") or "").strip()
        file_type = str(item.get("type") or "").strip().lower()
        if file_type != "pdf" and "pdf" not in mime_type.lower() and not filename.lower().endswith(".pdf"):
            continue
        ranked_files.append((score_europepmc_repo_pdf_file(item), len(filename), index, item))

    ranked_files.sort(key=lambda entry: (-entry[0], entry[1], entry[2]))
    return [item for _, _, _, item in ranked_files]


def append_unique_url(candidates: List[str], seen: set, url: Optional[str]):
    if not url:
        return
    clean_url = str(url).strip()
    if not clean_url or clean_url in seen:
        return
    seen.add(clean_url)
    candidates.append(clean_url)


def build_europepmc_repo_pdf_candidates(
    pmcid: Optional[str], fulltext_repo_metadata: Optional[Dict[str, Any]]
) -> List[str]:
    """Build only the repository-backed PDF candidates used by the browser Open PDF flow."""
    normalized_pmcid = normalize_pmcid(pmcid)
    candidates: List[str] = []
    seen = set()
    repo_version = stringify_repo_version((fulltext_repo_metadata or {}).get("version", 1))

    for item in extract_ranked_europepmc_repo_pdf_files(fulltext_repo_metadata):
        filename = str(item.get("filename") or item.get("fileName") or "").strip()
        mime_type = str(item.get("mimeType") or "").strip()
        direct_url = item.get("url")
        if normalized_pmcid and filename:
            append_unique_url(
                candidates,
                seen,
                "https://europepmc.org/api/fulltextRepo?"
                + urlencode(
                    {
                        "pmcId": normalized_pmcid,
                        "type": "FILE",
                        "fileName": filename,
                        "mimeType": mime_type or "application/pdf",
                        "version": repo_version,
                        "pmc_pageType": "pdf",
                        "pmc_domain": "null",
                    }
                ),
            )
        append_unique_url(candidates, seen, direct_url)

    return candidates


async def fetch_europepmc_record(
    pmid: Optional[str],
    pmcid: Optional[str],
    log_callback=None,
    client: Optional[httpx.AsyncClient] = None,
) -> Optional[Dict[str, Any]]:
    """通过 EuropePMC REST 获取文章元数据，用于解析更稳妥的 PDF 下载地址。"""
    queries = []
    if pmid:
        queries.append(f"EXT_ID:{pmid} AND SRC:MED")
    normalized_pmcid = normalize_pmcid(pmcid)
    if normalized_pmcid:
        queries.append(f"PMCID:{normalized_pmcid}")

    if not queries:
        return None

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json",
    }

    owned_client = client is None
    if owned_client:
        client = httpx.AsyncClient(timeout=20.0, verify=False, follow_redirects=True)

    try:
        for query in queries:
            try:
                if log_callback:
                    log_callback(f"PMID {pmid or 'N/A'}: EuropePMC metadata query -> {query}")
                resp = await client.get(
                    "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
                    params={"query": query, "format": "json", "resultType": "core", "pageSize": 1},
                    headers=headers,
                    timeout=20.0,
                )
                resp.raise_for_status()
                results = resp.json().get("resultList", {}).get("result", [])
                if results:
                    record = results[0]
                    if log_callback:
                        log_callback(
                            f"PMID {pmid or record.get('pmid', 'N/A')}: EuropePMC resolved "
                            f"pmcid={record.get('pmcid', 'N/A')} hasPDF={record.get('hasPDF', 'N/A')} "
                            f"inPMC={record.get('inPMC', 'N/A')}"
                        )
                    return record
                if log_callback:
                    log_callback(f"PMID {pmid or 'N/A'}: EuropePMC metadata query returned no results")
            except Exception as e:
                if log_callback:
                    log_callback(f"PMID {pmid or 'N/A'}: EuropePMC metadata query failed - {type(e).__name__}: {e}")
        return None
    finally:
        if owned_client and client is not None:
            await client.aclose()


async def fetch_europepmc_fulltext_repo_metadata(
    pmcid: Optional[str],
    pmid: Optional[str] = None,
    log_callback=None,
    client: Optional[httpx.AsyncClient] = None,
) -> Optional[Dict[str, Any]]:
    """Query the EuropePMC fulltext repository metadata used by the browser Open PDF flow."""
    normalized_pmcid = normalize_pmcid(pmcid)
    if not normalized_pmcid:
        return None

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json",
    }

    owned_client = client is None
    if owned_client:
        client = httpx.AsyncClient(timeout=20.0, verify=False, follow_redirects=True)

    try:
        if log_callback:
            log_callback(f"PMID {pmid or 'N/A'}: EuropePMC fulltextRepo metadata query -> {normalized_pmcid}")

        resp = await client.get(
            "https://europepmc.org/api/fulltextRepo",
            params={"pmcId": normalized_pmcid, "type": "METADATA"},
            headers=headers,
            timeout=20.0,
        )
        resp.raise_for_status()
        data = resp.json()

        if not data.get("success", True):
            if log_callback:
                log_callback(f"PMID {pmid or 'N/A'}: EuropePMC fulltextRepo metadata returned success=false")
            return None

        pdf_names = [
            str(item.get("filename") or item.get("fileName") or "").strip()
            for item in extract_ranked_europepmc_repo_pdf_files(data)
        ]

        if log_callback:
            if pdf_names:
                preview = ", ".join(pdf_names[:3])
                suffix = " ..." if len(pdf_names) > 3 else ""
                log_callback(
                    f"PMID {pmid or data.get('pmid', 'N/A')}: EuropePMC fulltextRepo found "
                    f"{len(pdf_names)} PDF file(s): {preview}{suffix}"
                )
            else:
                log_callback(f"PMID {pmid or data.get('pmid', 'N/A')}: EuropePMC fulltextRepo found no PDF files")

        return data
    except Exception as e:
        if log_callback:
            log_callback(f"PMID {pmid or 'N/A'}: EuropePMC fulltextRepo metadata query failed - {type(e).__name__}: {e}")
        return None
    finally:
        if owned_client and client is not None:
            await client.aclose()


def build_europepmc_pdf_candidates(
    paper: Dict,
    epmc_record: Optional[Dict[str, Any]],
    fulltext_repo_metadata: Optional[Dict[str, Any]] = None,
) -> Tuple[Optional[str], List[str]]:
    """构造 EuropePMC PDF 候选链接，优先使用真实浏览器走的 fulltextRepo 链路。"""
    pmcid = normalize_pmcid(paper.get("pmcid"))
    if epmc_record and epmc_record.get("pmcid"):
        pmcid = normalize_pmcid(epmc_record["pmcid"])

    candidates: List[str] = []
    seen = set()

    for url in build_europepmc_repo_pdf_candidates(pmcid, fulltext_repo_metadata):
        append_unique_url(candidates, seen, url)

    fulltext_urls = []
    if epmc_record:
        fulltext_urls = epmc_record.get("fullTextUrlList", {}).get("fullTextUrl", [])
        if isinstance(fulltext_urls, dict):
            fulltext_urls = [fulltext_urls]

    europepmc_pdf_urls = []
    external_pdf_urls = []
    for item in fulltext_urls:
        if not isinstance(item, dict):
            continue
        if item.get("availabilityCode") != "OA" or item.get("documentStyle") != "pdf":
            continue
        if item.get("site") == "Europe_PMC":
            europepmc_pdf_urls.append(item.get("url"))
        else:
            external_pdf_urls.append(item.get("url"))

    for url in europepmc_pdf_urls:
        append_unique_url(candidates, seen, url)
    for url in external_pdf_urls:
        append_unique_url(candidates, seen, url)
    if pmcid:
        append_unique_url(candidates, seen, f"https://europepmc.org/articles/{pmcid}?pdf=render")
        append_unique_url(candidates, seen, f"https://europepmc.org/backend/ptpmcrender.fcgi?accid={pmcid}&blobtype=pdf")

    return pmcid or None, candidates


async def try_europepmc_candidate_urls(
    pmid: str,
    candidate_urls: List[str],
    pdf_path: str,
    md_path: str,
    client: httpx.AsyncClient,
    log_callback,
    tried_urls: Optional[set] = None,
) -> bool:
    """Attempt candidate URLs in order and stop on the first valid PDF."""
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    max_attempts = 2
    tried_urls = tried_urls if tried_urls is not None else set()
    pending_urls = [url for url in candidate_urls if url not in tried_urls]

    for url_index, url in enumerate(pending_urls, start=1):
        tried_urls.add(url)
        for attempt in range(1, max_attempts + 1):
            if log_callback:
                log_callback(
                    f"PMID {pmid}: EuropePMC try {url_index}/{len(pending_urls)} "
                    f"(attempt {attempt}/{max_attempts}) -> {url}"
                )

            total_bytes = 0
            first_bytes = b""
            start_time = time.time()

            try:
                async with client.stream("GET", url, headers=headers, timeout=60.0, follow_redirects=True) as response:
                    final_url = str(response.url)
                    content_type = response.headers.get("content-type", "")
                    content_length = response.headers.get("content-length", "unknown")
                    if log_callback:
                        log_callback(
                            f"PMID {pmid}: response status={response.status_code} "
                            f"type={content_type or 'N/A'} len={content_length} final={final_url}"
                        )

                    if response.status_code != 200:
                        continue

                    with open(pdf_path, "wb") as f:
                        async for chunk in response.aiter_bytes():
                            if not first_bytes and chunk:
                                first_bytes = chunk[:16]
                            f.write(chunk)
                            total_bytes += len(chunk)

                elapsed = time.time() - start_time
                pdf_signature = b"%PDF" in first_bytes[:16]
                looks_like_pdf = (
                    "pdf" in content_type.lower()
                    or pdf_signature
                    or final_url.lower().endswith(".pdf")
                    or "pdf=render" in final_url.lower()
                    or "fulltextrepo" in final_url.lower()
                )

                if total_bytes < 1000:
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                    if log_callback:
                        log_callback(f"PMID {pmid}: downloaded file too small ({total_bytes} bytes), retrying next URL")
                    continue

                if not looks_like_pdf:
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                    if log_callback:
                        prefix = first_bytes[:16].hex() if first_bytes else "empty"
                        log_callback(
                            f"PMID {pmid}: response does not look like PDF "
                            f"(type={content_type or 'N/A'}, first16={prefix}), retrying next URL"
                        )
                    continue

                kb_size = total_bytes / 1024
                speed = kb_size / elapsed if elapsed > 0 else 0
                if os.path.exists(md_path):
                    try:
                        os.remove(md_path)
                        if log_callback:
                            log_callback(f"PMID {pmid}: removed stale fallback MD after successful PMC download")
                    except Exception:
                        pass
                if log_callback:
                    log_callback(
                        f"PMID {pmid}: PDF (PMC) | {kb_size:.0f}KB | {speed:.0f}KB/s | {elapsed:.1f}s | {final_url}"
                    )
                return True

            except Exception as e:
                if os.path.exists(pdf_path):
                    try:
                        os.remove(pdf_path)
                    except Exception:
                        pass
                if log_callback:
                    log_callback(
                        f"PMID {pmid}: EuropePMC request failed on attempt {attempt}/{max_attempts} "
                        f"- {type(e).__name__}: {e}"
                    )

    return False


async def download_from_europepmc(paper: Dict, save_dir: str, log_callback) -> bool:
    """从 EuropePMC 下载 PMC 文献 PDF，优先走 REST 解析到的官方 PDF 地址。"""
    pmid = paper["pmid"]
    pdf_path = os.path.join(save_dir, f"{pmid}.pdf")
    md_path = os.path.join(save_dir, f"{pmid}.md")

    if os.path.exists(pdf_path):
        if log_callback:
            log_callback(f"PMID {pmid}: PMC download skipped, existing PDF found")
        return True

    source_pmcid = normalize_pmcid(paper.get("pmcid"))
    if log_callback:
        log_callback(f"PMID {pmid}: starting EuropePMC download (xml_pmcid={source_pmcid or 'N/A'})")

    fulltext_repo_metadata = None
    epmc_record = None
    resolved_pmcid = source_pmcid
    tried_urls = set()

    async with httpx.AsyncClient(timeout=60.0, verify=False, follow_redirects=True) as client:
        if source_pmcid:
            fulltext_repo_metadata = await fetch_europepmc_fulltext_repo_metadata(
                source_pmcid, pmid, log_callback, client=client
            )
            fast_candidates = build_europepmc_repo_pdf_candidates(source_pmcid, fulltext_repo_metadata)
            if log_callback:
                log_callback(f"PMID {pmid}: resolved EuropePMC pmcid={source_pmcid}")
            if fast_candidates:
                if await try_europepmc_candidate_urls(
                    pmid,
                    fast_candidates,
                    pdf_path,
                    md_path,
                    client,
                    log_callback,
                    tried_urls=tried_urls,
                ):
                    return True
                if log_callback:
                    log_callback(f"PMID {pmid}: fulltextRepo fast path exhausted, loading REST fallback")
            elif log_callback:
                log_callback(f"PMID {pmid}: no fast PMC candidates from fulltextRepo, loading REST fallback")

        epmc_record = await fetch_europepmc_record(pmid, source_pmcid, log_callback, client=client)
        resolved_pmcid = normalize_pmcid((epmc_record or {}).get("pmcid")) or source_pmcid
        if resolved_pmcid and (fulltext_repo_metadata is None or resolved_pmcid != source_pmcid):
            fulltext_repo_metadata = await fetch_europepmc_fulltext_repo_metadata(
                resolved_pmcid, pmid, log_callback, client=client
            )

        resolved_pmcid, candidate_urls = build_europepmc_pdf_candidates(paper, epmc_record, fulltext_repo_metadata)

        if log_callback:
            log_callback(f"PMID {pmid}: resolved EuropePMC pmcid={resolved_pmcid or 'N/A'}")

        if not candidate_urls:
            if log_callback:
                log_callback(f"PMID {pmid}: no EuropePMC PDF candidates found")
            return False

        if await try_europepmc_candidate_urls(
            pmid,
            candidate_urls,
            pdf_path,
            md_path,
            client,
            log_callback,
            tried_urls=tried_urls,
        ):
            return True

    if log_callback:
        log_callback(f"PMID {pmid}: all EuropePMC candidates failed")
    return False


async def download_scihub_with_mirrors(
    doi: str,
    save_path: str,
    mirrors: List[str],
    log_callback
) -> Tuple[Optional[str], float, float]:
    """
    按指定镜像顺序尝试下载，返回下载速度信息
    返回: (结果描述, 文件大小KB, 下载速度KB/s)
    """
    if not HAS_CURL_CFFI:
        return None, 0, 0

    pmid = os.path.basename(save_path).replace('.pdf', '')

    async with AsyncSession(impersonate="chrome120", verify=False) as browser:
        for i, mirror in enumerate(mirrors):
            mirror_name = mirror.split('.')[-1]
            try:
                page_url = f"{mirror}/{doi}"
                start_time = time.time()

                try:
                    resp = await browser.get(page_url, timeout=15, allow_redirects=True)
                except Exception as e:
                    if log_callback:
                        log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): Connection failed - {str(e)[:30]}")
                    continue

                if resp.status_code != 200:
                    if log_callback:
                        log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): HTTP {resp.status_code}")
                    continue

                content_type = resp.headers.get('Content-Type', '').lower()
                if 'application/pdf' in content_type and len(resp.content) > 1000:
                    elapsed = time.time() - start_time
                    kb_size = len(resp.content) / 1024
                    speed = kb_size / elapsed if elapsed > 0 else 0
                    with open(save_path, "wb") as f:
                        f.write(resp.content)
                    md_path = save_path.replace(".pdf", ".md")
                    if os.path.exists(md_path):
                        try:
                            os.remove(md_path)
                        except Exception:
                            pass
                    return f"PDF (Sci-Hub {mirror_name}) | {kb_size:.0f}KB | {speed:.0f}KB/s | {elapsed:.1f}s", kb_size, speed

                soup = BeautifulSoup(resp.text, 'lxml')
                pdf_url = None

                pdf_div = soup.find('div', class_='pdf')
                if pdf_div:
                    obj_tag = pdf_div.find('object', attrs={'type': 'application/pdf'})
                    if obj_tag and obj_tag.get('data'):
                        pdf_url = obj_tag['data']

                if not pdf_url:
                    download_div = soup.find('div', class_='download')
                    if download_div:
                        link = download_div.find('a', href=True)
                        if link:
                            pdf_url = link['href']

                if not pdf_url:
                    embed = soup.find('embed', attrs={'type': 'application/pdf'})
                    if embed and embed.get('src'):
                        pdf_url = embed['src']

                if not pdf_url:
                    iframe = soup.find('iframe', id='pdf')
                    if iframe and iframe.get('src'):
                        pdf_url = iframe['src']

                if not pdf_url:
                    # 检查是否触发反爬
                    if "captcha" in resp.text.lower() or "checking your browser" in resp.text.lower():
                        if log_callback:
                            log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): CAPTCHA triggered")
                    else:
                        if log_callback:
                            log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): No PDF link found")
                    continue

                if '#' in pdf_url:
                    pdf_url = pdf_url.split('#')[0]
                if pdf_url.startswith('//'):
                    pdf_url = 'https:' + pdf_url
                elif pdf_url.startswith('/'):
                    pdf_url = mirror + pdf_url
                elif not pdf_url.startswith('http'):
                    pdf_url = mirror + '/' + pdf_url

                download_start = time.time()
                try:
                    pdf_resp = await browser.get(pdf_url, timeout=120, allow_redirects=True)
                except Exception as e:
                    if log_callback:
                        log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): PDF download failed - {str(e)[:30]}")
                    continue

                if pdf_resp.status_code == 200 and len(pdf_resp.content) > 1000:
                    elapsed = time.time() - download_start
                    kb_size = len(pdf_resp.content) / 1024
                    speed = kb_size / elapsed if elapsed > 0 else 0
                    with open(save_path, "wb") as f:
                        f.write(pdf_resp.content)
                    md_path = save_path.replace(".pdf", ".md")
                    if os.path.exists(md_path):
                        try:
                            os.remove(md_path)
                        except Exception:
                            pass
                    return f"PDF (Sci-Hub {mirror_name}) | {kb_size:.0f}KB | {speed:.0f}KB/s | {elapsed:.1f}s", kb_size, speed
                else:
                    if log_callback:
                        log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): PDF too small or bad status")
                    continue

            except Exception as e:
                if log_callback:
                    log_callback(f"  [{pmid}] Mirror {i+1}/{len(mirrors)} ({mirror_name}): Error - {str(e)[:30]}")
                continue

    return None, 0, 0


async def download_papers_concurrent(
    papers: List[Dict],
    save_dir: str,
    available_mirrors: List[str],
    log_callback
) -> Tuple[List[str], List[str]]:
    """
    并发下载论文，返回 (成功列表, 失败列表)

    并发数 = len(available_mirrors)
    每个 worker 使用不同的镜像顺序
    """
    success_pmids = []
    failed_pmids = []
    skipped_pmids = []

    # 检查已存在的文件
    papers_to_download = []
    for paper in papers:
        pmid = paper['pmid']
        pdf_path = os.path.join(save_dir, f"{pmid}.pdf")
        md_path = os.path.join(save_dir, f"{pmid}.md")
        if os.path.exists(pdf_path):
            skipped_pmids.append(pmid)
            log_callback(f"PMID {pmid}: Skipped (PDF exists)")
        elif os.path.exists(md_path) and not (paper.get('pmcid') or paper.get('doi')):
            skipped_pmids.append(pmid)
            log_callback(f"PMID {pmid}: Skipped (MD fallback exists, no PMC/DOI retry path)")
        else:
            if os.path.exists(md_path):
                log_callback(f"PMID {pmid}: Existing MD fallback found, retrying PDF download")
            papers_to_download.append(paper)

    if not papers_to_download:
        log_callback("All papers are already downloaded; no duplicate download is needed.")
        return skipped_pmids, []

    # 分离 PMC 文献和非 PMC 文献
    pmc_papers = [p for p in papers_to_download if p.get('pmcid')]
    non_pmc_papers = [p for p in papers_to_download if not p.get('pmcid')]

    log_callback(f"Papers pending download: {len(papers_to_download)} total, {len(pmc_papers)} PMC, {len(non_pmc_papers)} requiring Sci-Hub")

    # 1. 先处理 PMC 文献
    if pmc_papers:
        log_callback(f"Downloading PMC papers ({len(pmc_papers)})...")
        for paper in pmc_papers:
            result = await download_from_europepmc(paper, save_dir, log_callback)
            if result:
                success_pmids.append(paper['pmid'])
            else:
                # PMC 失败的加入非 PMC 队列尝试 Sci-Hub
                non_pmc_papers.append(paper)
                log_callback(f"PMID {paper['pmid']}: PMC download failed; queued for Sci-Hub fallback")

    # 2. 并发处理非 PMC 文献 (Sci-Hub)
    if not non_pmc_papers:
        log_callback("All PMC papers downloaded.")
        return success_pmids + skipped_pmids, failed_pmids

    if not available_mirrors:
        log_callback("No available mirrors; Sci-Hub papers cannot be downloaded.")
        for paper in non_pmc_papers:
            failed_pmids.append(paper['pmid'])
            # 生成 MD 备份
            md_path = os.path.join(save_dir, f"{paper['pmid']}.md")
            content = f"# {paper.get('title')}\n\nPMID: {paper['pmid']}\nDOI: {paper.get('doi')}\n\n---\n\n{paper.get('abstract')}"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content)
        return success_pmids + skipped_pmids, failed_pmids

    num_workers = len(available_mirrors)
    paper_queue = asyncio.Queue()

    # 填充队列
    for paper in non_pmc_papers:
        await paper_queue.put(paper)

    log_callback(f"Starting {num_workers} concurrent workers for Sci-Hub downloads...")

    # 用于线程安全的结果收集
    results_lock = asyncio.Lock()

    async def worker(worker_id: int):
        """每个 worker 使用自己的镜像顺序"""
        rotated_mirrors = create_rotated_mirror_order(available_mirrors, worker_id)

        while True:
            try:
                paper = await asyncio.wait_for(paper_queue.get(), timeout=1.0)
            except asyncio.TimeoutError:
                break  # 队列空了

            pmid = paper['pmid']
            doi = paper.get('doi')

            if not doi:
                async with results_lock:
                    failed_pmids.append(pmid)
                log_callback(f"PMID {pmid}: No DOI available; Sci-Hub download is not possible")
                # 生成 MD 备份
                md_path = os.path.join(save_dir, f"{pmid}.md")
                content = f"# {paper.get('title')}\n\nPMID: {pmid}\nDOI: N/A\n\n---\n\n{paper.get('abstract')}"
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(content)
                paper_queue.task_done()
                continue

            # 使用轮转后的镜像顺序尝试下载
            pdf_path = os.path.join(save_dir, f"{pmid}.pdf")
            result, kb_size, speed = await download_scihub_with_mirrors(
                doi,
                pdf_path,
                rotated_mirrors,
                log_callback
            )

            async with results_lock:
                if result:
                    success_pmids.append(pmid)
                    log_callback(f"PMID {pmid}: {result}")
                else:
                    failed_pmids.append(pmid)
                    log_callback(f"PMID {pmid}: Sci-Hub download failed")
                    # 生成 MD 备份
                    md_path = os.path.join(save_dir, f"{pmid}.md")
                    content = f"# {paper.get('title')}\n\nPMID: {pmid}\nDOI: {doi}\n\n---\n\n{paper.get('abstract')}"
                    with open(md_path, "w", encoding="utf-8") as f:
                        f.write(content)

            paper_queue.task_done()

    # 启动 worker
    workers = [asyncio.create_task(worker(i)) for i in range(num_workers)]
    await asyncio.gather(*workers)

    return success_pmids + skipped_pmids, failed_pmids


def save_failed_literature_list(save_dir: str, failed_pmids: List[str]):
    """保存下载失败的文献 PMID 列表"""
    if not failed_pmids:
        return
    output_path = os.path.join(save_dir, "literature_without_pdf.txt")
    with open(output_path, "w", encoding="utf-8") as f:
        for pmid in failed_pmids:
            f.write(f"{pmid}\n")


def create_literature_archive(save_dir: str, log_callback) -> Optional[str]:
    """将所有文献打包为 downloaded_literature.zip"""
    # 收集所有 PDF 和 MD 文件
    pdf_files = glob.glob(os.path.join(save_dir, "*.pdf"))
    md_files = glob.glob(os.path.join(save_dir, "*.md"))
    files = pdf_files + md_files

    if not files:
        log_callback("No files to archive")
        return None

    zip_path = os.path.join(save_dir, "downloaded_literature.zip")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            zf.write(f, os.path.basename(f))

    log_callback(f"Created archive: {zip_path} ({len(files)} files)")
    return zip_path

async def run_fulltext_download_pipeline(collection_name: str, log_callback) -> Tuple[List[str], List[str]]:
    save_path = get_collection_path(collection_name)
    xml_path = os.path.join(save_path, "pubmed_results.xml")
    if not os.path.exists(xml_path):
        log_callback("XML not found.")
        return [], []

    with open(xml_path, "rb") as f:
        arts = parse_pubmed_xml_details(f.read())

    if not arts:
        log_callback("No articles found in XML.")
        return [], []

    log_callback(f"Found {len(arts)} articles to download")
    log_callback("=" * 50)
    log_callback("Testing Sci-Hub mirrors (parallel)...")
    available_mirrors = await test_scihub_mirrors(log_callback)
    log_callback(f"Found {len(available_mirrors)} available mirrors")
    log_callback("=" * 50)

    success, failed = await download_papers_concurrent(
        arts, save_path, available_mirrors, log_callback
    )

    if failed:
        save_failed_literature_list(save_path, failed)
        log_callback(f"Saved {len(failed)} failed PMIDs to literature_without_pdf.txt")

    create_literature_archive(save_path, log_callback)
    log_callback(f"Download complete: {len(success)} success, {len(failed)} failed")
    return success, failed

# ==========================================
#               3. LLM & Vector Helpers
# ==========================================

async def fetch_llm_models(config: Dict) -> Tuple[bool, List[str], str]:
    """
    从 LLM API 拉取可用模型列表
    返回: (成功与否, 模型列表, 错误信息)
    """
    try:
        base_url = config.get('base_url', '').rstrip('/')
        if not base_url:
            return False, [], "Base URL is empty"

        if not base_url.endswith('/v1'):
            base_url += '/v1'

        url = f"{base_url}/models"
        headers = {
            "Authorization": f"Bearer {config.get('api_key', '')}",
            "Content-Type": "application/json"
        }

        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()
            data = resp.json()

            models = []
            if 'data' in data:
                for model in data['data']:
                    model_id = model.get('id', '')
                    if model_id:
                        models.append(model_id)

            models.sort()
            return True, models, ""
    except httpx.TimeoutException:
        return False, [], "Connection timeout"
    except httpx.HTTPStatusError as e:
        return False, [], f"HTTP Error: {e.response.status_code}"
    except Exception as e:
        return False, [], str(e)


async def test_llm_connection(config: Dict) -> Tuple[bool, str, float]:
    """
    测试 LLM 连接是否正常
    返回: (成功与否, 响应内容/错误信息, 响应时间秒)
    """
    try:
        base_url = config.get('base_url', '').rstrip('/')
        if not base_url:
            return False, "Base URL is empty", 0

        if not base_url.endswith('/v1'):
            base_url += '/v1'

        url = f"{base_url}/chat/completions"
        headers = {
            "Authorization": f"Bearer {config.get('api_key', '')}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": config.get('model', ''),
            "messages": [{"role": "user", "content": "Hello"}],
            "max_tokens": 10,
            "temperature": 0
        }

        start_time = time.time()
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(url, headers=headers, json=payload)
            resp.raise_for_status()
            elapsed = time.time() - start_time

            data = resp.json()
            if 'choices' in data and len(data['choices']) > 0:
                content = data['choices'][0].get('message', {}).get('content', '')
                return True, content[:100], elapsed
            return True, "Response received", elapsed
    except httpx.TimeoutException:
        return False, "Connection timeout", 0
    except httpx.HTTPStatusError as e:
        return False, f"HTTP Error: {e.response.status_code}", 0
    except Exception as e:
        return False, str(e), 0

def _parse_chat_completion_content(data: Dict[str, Any]) -> str:
    try:
        return data['choices'][0].get('message', {}).get('content', '') or ''
    except Exception:
        return ''

async def _call_llm_openai_nonstream(url: str, headers: Dict[str, str], payload: Dict[str, Any], timeout: httpx.Timeout, log_callback) -> Optional[str]:
    request_payload = dict(payload)
    request_payload["stream"] = False
    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(url, headers=headers, json=request_payload)
        resp.raise_for_status()
        return _parse_chat_completion_content(resp.json())

async def _call_llm_openai_stream(url: str, headers: Dict[str, str], payload: Dict[str, Any], timeout: httpx.Timeout, log_callback) -> Optional[str]:
    request_payload = dict(payload)
    request_payload["stream"] = True
    content_parts: List[str] = []
    first_content_seen = False
    start_time = time.time()

    async with httpx.AsyncClient(timeout=timeout) as client:
        async with client.stream("POST", url, headers=headers, json=request_payload) as resp:
            resp.raise_for_status()
            async for line in resp.aiter_lines():
                if not line:
                    continue
                line = line.strip()
                if line.startswith("data:"):
                    line = line[5:].strip()
                if not line or line == "[DONE]":
                    continue
                try:
                    chunk = json.loads(line)
                except json.JSONDecodeError:
                    continue

                choice = (chunk.get("choices") or [{}])[0]
                delta = choice.get("delta") or {}
                piece = delta.get("content")
                if piece is None:
                    # Some providers send a final non-delta message-shaped chunk.
                    piece = choice.get("message", {}).get("content")
                if piece:
                    if not first_content_seen:
                        first_content_seen = True
                        log_callback("  - LLM stream: first content chunk received.")
                    content_parts.append(piece)

    content = "".join(content_parts).strip()
    if content:
        log_callback(f"  - LLM stream completed in {time.time() - start_time:.1f}s.")
    return content

async def call_llm_openai(prompt: str, config: Dict, log_callback, system_prompt: Optional[str] = None) -> Optional[str]:
    headers = {"Authorization": f"Bearer {config['api_key']}", "Content-Type": "application/json"}
    base_url = config['base_url'].rstrip('/')
    if not base_url.endswith('/v1'): base_url += '/v1'
    url = f"{base_url}/chat/completions"
    use_stream = config.get("stream", True)
    request_timeout = float(config.get("timeout", 600))
    read_timeout = float(config.get("read_timeout", 120))
    timeout = httpx.Timeout(request_timeout, connect=30.0, read=read_timeout, write=60.0, pool=30.0)
    payload = {
        "model": config['model'],
        "messages": [
            {"role": "system", "content": system_prompt or "You are a professional academic assistant."},
            {"role": "user", "content": prompt},
        ],
        "temperature": config.get('temperature', 0.1),
    }
    try:
        if use_stream:
            log_callback("  - LLM request mode: stream.")
            try:
                content = await _call_llm_openai_stream(url, headers, payload, timeout, log_callback)
                if content:
                    return content
                log_callback("  - LLM stream returned empty content; retrying without stream.")
            except httpx.HTTPStatusError as stream_error:
                status = stream_error.response.status_code if stream_error.response is not None else "N/A"
                log_callback(f"  - LLM stream HTTP error {status}; retrying without stream.")
            except Exception as stream_error:
                log_callback(f"LLM Request Error: stream failed - {stream_error}")
                return None

        log_callback("  - LLM request mode: non-stream.")
        return await _call_llm_openai_nonstream(url, headers, payload, timeout, log_callback)
    except Exception as e:
        log_callback(f"LLM Request Error: {e}")
        return None

def validate_csv_content(csv_text: str, log_callback) -> Optional[pd.DataFrame]:
    match = re.search(r"```csv\n(.*?)```", csv_text, re.DOTALL)
    raw_csv = match.group(1).strip() if match else csv_text.strip()
    try:
        data = []
        lines = raw_csv.split('\n')
        serial_pattern = r'^\d+(\.\d+)*$'
        for line in lines:
            parts = line.strip().split(',', 1)
            if len(parts) < 2: continue
            serial = parts[0].strip().strip('"').strip("'")
            title = parts[1].strip().strip('"').strip("'")
            if re.match(serial_pattern, serial): data.append({'Serial': serial, 'Title': title})
        return pd.DataFrame(data)[['Serial', 'Title']] if data else None
    except: return None

async def get_embeddings_batch(texts: List[str], config: Dict, log_callback=None) -> List[List[float]]:
    if not texts: return []
    res = []
    headers = {"Authorization": f"Bearer {config['api_key']}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=60.0) as client:
        for i in range(0, len(texts), config.get('batch_size', 64)):
            batch = texts[i : i + config.get('batch_size', 64)]
            try:
                resp = await client.post(config['base_url'], headers=headers, json={"input": batch, "model": config['model_name']})
                if resp.status_code == 200: res.extend([x['embedding'] for x in resp.json()['data']])
                else: res.extend([None]*len(batch))
            except: res.extend([None]*len(batch))
    return res

async def search_text_from_chromadb(collection_name: str, query: str, k: int, config: Dict) -> List[Dict]:
    path = os.path.join(get_collection_path(collection_name), collection_name)
    if not os.path.exists(path): return []
    try:
        q_emb = await get_embeddings_batch([query], config)
        if not q_emb or not q_emb[0]: return []
        client = await asyncio.to_thread(chromadb.PersistentClient, path=path)
        coll = await asyncio.to_thread(client.get_collection, name=collection_name)
        res = await asyncio.to_thread(coll.query, query_embeddings=[q_emb[0]], n_results=k, include=['metadatas', 'documents'])
        ret = []
        if res['documents'] and res['documents'][0]:
            for i, doc in enumerate(res['documents'][0]):
                if res['ids'][0][i] in [ALL_ABSTRACTS_DOC_ID, TITLE_PMID_DOI_DOC_ID]: continue
                m = res['metadatas'][0][i]
                ret.append({"content": doc, "pmid": m.get('pmid', 'N/A'), "source": m.get('title', 'Unknown')})
        return ret
    except: return []

# ==========================================
#               4. Download & Process Helpers
# ==========================================

async def download_scihub_with_fallback(doi: str, save_path: str, normal_client: httpx.AsyncClient, log_callback=None, available_mirrors: List[str] = None) -> Optional[str]:
    def log(msg, level="info"):
        if level == "warning": logger.warning(msg)
        else: logger.info(msg)
        if log_callback: log_callback(f"  {msg}")

    if not HAS_CURL_CFFI:
        log("❌ curl_cffi module not found. Cannot bypass Sci-Hub protection.", "warning")
        return None

    mirrors_to_try = available_mirrors if available_mirrors else SCIHUB_MIRRORS
    
    async with AsyncSession(impersonate="chrome120", verify=False) as browser:
        for i, mirror in enumerate(mirrors_to_try):
            try:
                page_url = f"{mirror}/{doi}"
                log(f"-> [Mirror {i+1}] Connecting: {mirror}")
                
                t0 = time.time()
                try:
                    resp = await browser.get(page_url, timeout=15, allow_redirects=True)
                except Exception as e:
                    log(f"   Connection failed: {e} ({time.time()-t0:.1f}s)", "warning")
                    continue
                
                if resp.status_code != 200:
                    log(f"   Non-200 status code: {resp.status_code}", "warning")
                    continue

                content_type = resp.headers.get('Content-Type', '').lower()
                if 'application/pdf' in content_type:
                    log("   PDF returned directly")
                    with open(save_path, "wb") as f: f.write(resp.content)
                    md_path = save_path.replace(".pdf", ".md")
                    if os.path.exists(md_path):
                        try:
                            os.remove(md_path)
                        except Exception:
                            pass
                    return f"PDF (Sci-Hub Direct {mirror})"

                soup = BeautifulSoup(resp.text, 'lxml')
                pdf_url = None

                pdf_div = soup.find('div', class_='pdf')
                if pdf_div:
                    obj_tag = pdf_div.find('object', attrs={'type': 'application/pdf'})
                    if obj_tag and obj_tag.get('data'): pdf_url = obj_tag['data']

                if not pdf_url:
                    download_div = soup.find('div', class_='download')
                    if download_div:
                        link = download_div.find('a', href=True)
                        if link: pdf_url = link['href']

                if not pdf_url:
                    embed = soup.find('embed', attrs={'type': 'application/pdf'})
                    if embed and embed.get('src'): pdf_url = embed['src']
                
                if not pdf_url:
                    iframe = soup.find('iframe', id='pdf')
                    if iframe and iframe.get('src'): pdf_url = iframe['src']

                if not pdf_url:
                    button = soup.find('button', onclick=lambda x: x and 'location.href=' in x)
                    if button:
                        match = re.search(r"location\.href='([^']+)'", button['onclick'])
                        if match: pdf_url = match.group(1)
                
                if not pdf_url:
                    if "captcha" in resp.text.lower() or "checking your browser" in resp.text.lower():
                        log("   CAPTCHA / anti-bot challenge triggered", "warning")
                    else:
                        log("   PDF link not found on page", "warning")
                    continue

                if '#' in pdf_url: pdf_url = pdf_url.split('#')[0]
                if pdf_url.startswith('//'): pdf_url = 'https:' + pdf_url
                elif pdf_url.startswith('/'): pdf_url = mirror + pdf_url
                elif not pdf_url.startswith('http'): pdf_url = mirror + '/' + pdf_url

                t2 = time.time()
                pdf_resp = await browser.get(pdf_url, timeout=120, allow_redirects=True)
                t_download = time.time() - t2
                
                if pdf_resp.status_code == 200:
                    content_len = len(pdf_resp.content)
                    kb_size = content_len / 1024
                    speed = kb_size / t_download if t_download > 0 else 0
                    
                    log(f"   Download succeeded: {kb_size:.0f}KB | {speed:.0f}KB/s | {t_download:.1f}s")
                    
                    if content_len > 1000:
                        with open(save_path, "wb") as f: f.write(pdf_resp.content)
                        md_path = save_path.replace(".pdf", ".md")
                        if os.path.exists(md_path):
                            try:
                                os.remove(md_path)
                            except Exception:
                                pass
                        return f"PDF (Sci-Hub {mirror.split('.')[-1]} - {speed:.0f}KB/s)"
                    else:
                        log("   File too small (<1KB)", "warning")
                        if os.path.exists(save_path): os.remove(save_path)
                else:
                    log(f"   PDF request failed: {pdf_resp.status_code}", "warning")

            except Exception as e:
                log(f"   Unknown error: {e}", "warning")
                continue
    return None

async def process_paper_download(paper: Dict, save_dir: str, log_callback=None):
    global _MIRRORS_TESTED, _AVAILABLE_MIRRORS_CACHE
    pmid = paper['pmid']
    pdf_path = os.path.join(save_dir, f"{pmid}.pdf")
    md_path = os.path.join(save_dir, f"{pmid}.md")

    if os.path.exists(pdf_path):
        return "Skipped (PDF exists)"
    if os.path.exists(md_path) and not (paper.get('pmcid') or paper.get('doi')):
        return "Skipped (MD exists, no retry path)"

    if not _MIRRORS_TESTED:
        log_callback("=" * 50)
        log_callback("Testing Sci-Hub mirror availability...")
        log_callback("=" * 50)
        _AVAILABLE_MIRRORS_CACHE = await test_scihub_mirrors(log_callback)
        _MIRRORS_TESTED = True
        log_callback("=" * 50)
    
    async with httpx.AsyncClient(timeout=60.0, verify=False) as client:
        if paper.get('pmcid'):
            if await download_from_europepmc(paper, save_dir, log_callback):
                return "PDF (EuropePMC)"
        
        if paper.get('doi'):
            source = await download_scihub_with_fallback(
                paper['doi'], pdf_path, client, 
                log_callback=log_callback, 
                available_mirrors=_AVAILABLE_MIRRORS_CACHE
            )
            if source:
                return source
        
        content = f"# {paper.get('title')}\n\nPMID: {pmid}\nDOI: {paper.get('doi')}\n\n---\n\n{paper.get('abstract')}"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(content)
        return "Abstract MD (Sci-Hub Failed)"

async def run_marker_conversion(content_dir: str, workers: int, log_callback, marker_batch_size: int = 10):
    if is_bundled_app():
        converter_executable = get_same_directory_converter_executable()
        if not converter_executable or not os.path.exists(converter_executable):
            log_callback("Error: bundled same-directory converter not found.")
            return
        cmd = [
            converter_executable,
            content_dir,
            "--pdftext_workers",
            str(workers),
            "--marker_batch_size",
            str(marker_batch_size),
        ]
    else:
        script = get_app_resource_path("convert_same_directory.py")
        if not os.path.exists(script):
            log_callback("Error: convert_same_directory.py not found.")
            return
        cmd = [
            sys.executable,
            "-u",
            script,
            content_dir,
            "--pdftext_workers",
            str(workers),
            "--marker_batch_size",
            str(marker_batch_size),
        ]

    def decode_process_output(raw: bytes) -> str:
        encodings = [
            "utf-8",
            locale.getpreferredencoding(False),
            sys.getfilesystemencoding(),
            "gb18030",
            "latin-1",
        ]
        seen = set()
        for enc in encodings:
            if not enc or enc.lower() in seen:
                continue
            seen.add(enc.lower())
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    log_callback(f"Running: {' '.join(cmd)}")
    try:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        env["PYTHONUTF8"] = "1"
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=env,
        )

        async def stream_pipe(pipe, prefix=""):
            buffer = bytearray()

            def emit(raw: bytes):
                msg = decode_process_output(raw).strip()
                if msg:
                    log_callback(f"{prefix}{msg}")

            while True:
                chunk = await pipe.read(4096)
                if not chunk:
                    break
                buffer.extend(chunk)

                while True:
                    newline_positions = [pos for pos in (buffer.find(b"\n"), buffer.find(b"\r")) if pos != -1]
                    if not newline_positions:
                        break
                    pos = min(newline_positions)
                    raw_line = bytes(buffer[:pos])
                    del buffer[:pos + 1]
                    if raw_line.strip():
                        emit(raw_line)

            if buffer.strip():
                emit(bytes(buffer))

        await asyncio.gather(
            stream_pipe(proc.stdout),
            stream_pipe(proc.stderr, "stderr: "),
        )
        return_code = await proc.wait()
        if return_code == 0:
            log_callback("Markdown conversion finished.")
        else:
            log_callback(f"Marker converter exited with code {return_code}.")
    except Exception as e:
        log_callback(f"Marker failed: {e}")

def preprocess_and_overwrite_md(file_path: str, pmid: str, log_callback) -> Optional[str]:
    try:
        with open(file_path, 'r', encoding='utf-8') as f: content = f.read()
        if not content.strip(): return None
        paras = [p.strip() for p in re.split(r'\n\s*\n', content.strip()) if p.strip()]
        filtered = [p for p in paras if not (p.startswith("- ") or (len(p) < 80 and not p.startswith("#")))]
        new_c = "\n\n".join(filtered)
        with open(file_path, 'w', encoding='utf-8') as f: f.write(new_c)
        return new_c
    except: return None

def paragraph_based_split_text(text: str) -> List[str]:
    chunks = []
    paras = [p.strip() for p in re.split(r'\n\s*\n', text.strip()) if p.strip()]
    for p in paras:
        if len(p) <= CONTENT_CHUNK_SIZE:
            if len(p) >= MIN_CHUNK_CHARS: chunks.append(p)
        else:
            for i in range(0, len(p), CONTENT_CHUNK_SIZE):
                sub = p[i:i+CONTENT_CHUNK_SIZE]
                if len(sub) >= MIN_CHUNK_CHARS: chunks.append(sub)
    return chunks

async def clear_vector_db(collection_name: str, log_callback):
    path = os.path.join(get_collection_path(collection_name), collection_name)
    if os.path.exists(path):
        try: shutil.rmtree(path); log_callback("√ Database cleared.")
        except Exception as e: log_callback(f"× Failed: {e}")

async def generate_debug_file(collection_name: str, log_callback):
    path = os.path.join(get_collection_path(collection_name), collection_name)
    if not os.path.exists(path): return
    try:
        client = await asyncio.to_thread(chromadb.PersistentClient, path=path)
        coll = await asyncio.to_thread(client.get_collection, name=collection_name)
        data = await asyncio.to_thread(coll.get, limit=5)
        with open(os.path.join(get_collection_path(collection_name), "test.txt"), "w") as f: f.write(str(data))
        log_callback("√ Debug file generated.")
    except: pass

async def process_vectorization(collection_name: str, config: Dict, log_callback):
    c_dir = get_collection_path(collection_name)
    db_path = os.path.join(c_dir, collection_name)
    os.makedirs(c_dir, exist_ok=True)
    xmls = glob.glob(os.path.join(c_dir, "*.xml"))
    meta_list = []
    if xmls:
        try:
            with open(xmls[0], "rb") as f: meta_list = parse_pubmed_xml_details(f.read())
        except: pass
    try:
        client = await asyncio.to_thread(chromadb.PersistentClient, path=db_path)
        coll = await asyncio.to_thread(client.get_or_create_collection, name=collection_name, metadata={"hnsw:space": "cosine"})
    except: return
    if meta_list:
        abs_json = json.dumps([{"pmid": p['pmid'], "abstract": p.get('abstract')} for p in meta_list if p.get('pmid')], ensure_ascii=False)
        meta_json = json.dumps([{"title": p.get('title'), "pmid": p.get('pmid'), "doi": p.get('doi')} for p in meta_list if p.get('pmid')], ensure_ascii=False)
        emb = await get_embeddings_batch(["placeholder"], config)
        if emb and emb[0]: coll.upsert(ids=[ALL_ABSTRACTS_DOC_ID, TITLE_PMID_DOI_DOC_ID], embeddings=[emb[0], emb[0]], documents=["META", "META"], metadatas=[{COMBINED_ABSTRACT_METADATA_KEY: abs_json}, {TITLE_PMID_DOI_METADATA_KEY: meta_json}])
    mds = glob.glob(os.path.join(c_dir, "*.md"))
    log_callback(f"Processing {len(mds)} MD files...")
    count = 0
    for md in mds:
        pmid = os.path.basename(md).replace(".md", "")
        txt = preprocess_and_overwrite_md(md, pmid, log_callback)
        if not txt: continue
        chunks = paragraph_based_split_text(txt)
        if not chunks: continue
        embs = await get_embeddings_batch(chunks, config)
        meta = next((p for p in meta_list if p['pmid'] == pmid), {"pmid": pmid})
        ids, embeddings, metadatas, docs = [], [], [], []
        for i, (chunk, emb) in enumerate(zip(chunks, embs)):
            if emb:
                ids.append(f"{pmid}_{i}"); embeddings.append(emb); docs.append(chunk)
                metadatas.append({"pmid": pmid, "title": meta.get("title", "") or "", "doi": meta.get("doi", "") or "", "chunk_index": i})
        if ids: coll.upsert(ids=ids, embeddings=embeddings, metadatas=metadatas, documents=docs); count += len(ids)
    log_callback(f"Vectorization done. Stored {count} chunks.")

# ==========================================
#               5. Main Logic Functions
# ==========================================

async def search_pubmed_and_save_xml(keyword, retmax, sort_by, min_date, max_date, save_dir, save_filename=None):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    search_params = {'db': 'pubmed', 'term': keyword, 'retmax': retmax, 'retmode': 'json', 'sort': 'pub date' if sort_by == 'pub_date' else 'relevance'}
    if min_date: search_params['mindate'] = min_date
    if max_date: search_params['maxdate'] = max_date
    
    resp = await asyncio.to_thread(make_request_with_retry, f"{base_url}/esearch.fcgi", search_params)
    id_list = resp.json().get('esearchresult', {}).get('idlist', [])
    if not id_list: return []

    fetch_params = {'db': 'pubmed', 'id': ",".join(id_list), 'retmode': 'xml'}
    resp_fetch = await asyncio.to_thread(make_request_with_retry, f"{base_url}/efetch.fcgi", fetch_params)
    
    if not save_filename: save_filename = "pubmed_results.xml"
    with open(os.path.join(save_dir, save_filename), "wb") as f: f.write(resp_fetch.content)
    return parse_pubmed_xml_details(resp_fetch.content)

async def generate_review_framework(collection_name: str, topic: str, custom_instructions: str, llm_config: Dict, log_callback, lang='zh'):
    c_dir = get_collection_path(collection_name)
    xml = os.path.join(c_dir, "pubmed_results.xml")
    if not os.path.exists(xml):
        log_callback("Error: pubmed_results.xml not found.")
        return

    log_callback("Step 1/4: Reading abstracts...")
    try:
        with open(xml, "rb") as f: arts = parse_pubmed_xml_details(f.read())
        if not arts: return
        abst_txt = "\n---\n".join([f"Title: {a.get('title')}\nAbstract: {a.get('abstract')}" for a in arts if a.get('abstract')])
    except: return

    log_callback("Step 2/4: Building prompt...")
    if lang == 'en':
        intro, base, fmt = FRAMEWORK_PROMPT_INTRO_EN, FRAMEWORK_PROMPT_BASE_REQS_EN, FRAMEWORK_PROMPT_OUTPUT_FMT_EN
    else:
        intro, base, fmt = FRAMEWORK_PROMPT_INTRO_ZH, FRAMEWORK_PROMPT_BASE_REQS_ZH, FRAMEWORK_PROMPT_OUTPUT_FMT_ZH
    
    full = f"{intro}\n{base}\n- {custom_instructions}\n{fmt}".format(topic=topic, abstracts=abst_txt)
    with open(os.path.join(c_dir, "prompt.txt"), "w", encoding="utf-8") as f: f.write(full)

    log_callback("Step 3/4: Requesting LLM...")
    res = await call_llm_openai(full, llm_config, log_callback)
    if not res: return

    log_callback("Step 4/4: Parsing CSV...")
    df = validate_csv_content(res, log_callback)
    if df is not None:
        out_csv = get_framework_csv_path(collection_name, lang)
        df.to_csv(out_csv, index=False, encoding="utf-8-sig")
        log_callback(f"√ Framework generated: {out_csv}")
    else: log_callback("× CSV parsing failed.")

def _section_sort_key(serial: str) -> List[int]:
    return [int(part) if str(part).isdigit() else 0 for part in str(serial).split('.')]

def _read_framework_rows(csv_path: str) -> List[Dict[str, str]]:
    df = pd.read_csv(csv_path, dtype=str)
    df.columns = ['Serial', 'Title'] + list(df.columns[2:])
    rows = []
    for _, row in df.iterrows():
        serial = str(row.get('Serial', '')).strip()
        title = str(row.get('Title', '')).strip()
        if serial:
            rows.append({"serial": serial, "title": title})
    return sorted(rows, key=lambda x: _section_sort_key(x["serial"]))

def parse_framework_to_sections(csv_path: str) -> Dict[str, Any]:
    """Build deterministic writing units from the framework tree.

    A natural writing unit is either:
    - an internal node whose direct children are all leaves, or
    - a leaf whose parent has mixed child depths, such as 5.5 beside 5.1-5.4.
    """
    try:
        rows = _read_framework_rows(csv_path)
        title_map = {r["serial"]: r["title"] for r in rows}
        serials = [r["serial"] for r in rows]
        serial_set = set(serials)
        children: Dict[str, List[str]] = {s: [] for s in serials}

        for serial in serials:
            if "." not in serial:
                continue
            parent = serial.rsplit(".", 1)[0]
            if parent in serial_set:
                children.setdefault(parent, []).append(serial)

        for child_ids in children.values():
            child_ids.sort(key=_section_sort_key)

        def is_leaf(serial: str) -> bool:
            return not children.get(serial)

        def parent_of(serial: str) -> Optional[str]:
            if "." not in serial:
                return None
            parent = serial.rsplit(".", 1)[0]
            return parent if parent in serial_set else None

        tasks: Dict[str, Any] = {}
        for serial in serials:
            child_ids = children.get(serial, [])
            parent = parent_of(serial)
            parent_children = children.get(parent, []) if parent else []
            parent_is_mixed = bool(parent_children) and any(not is_leaf(c) for c in parent_children) and any(is_leaf(c) for c in parent_children)

            if child_ids and all(is_leaf(c) for c in child_ids):
                unit_id = serial
                leaf_ids = child_ids
            elif is_leaf(serial) and (parent is None or parent_is_mixed):
                unit_id = serial
                leaf_ids = [serial]
            else:
                continue

            skeleton_ids = [unit_id] + [leaf_id for leaf_id in leaf_ids if leaf_id != unit_id]
            sibling_ids = [s for s in parent_children if s != unit_id] if parent else []
            tasks[unit_id] = {
                "id": unit_id,
                "title": title_map.get(unit_id, "Section"),
                "parent_id": parent,
                "parent_title": title_map.get(parent, "") if parent else "",
                "sibling_nodes": [{"serial": s, "title": title_map.get(s, "")} for s in sibling_ids],
                "skeleton": [f"{s} {title_map.get(s, '')}".strip() for s in skeleton_ids],
                "leaf_nodes": [{"serial": s, "title": title_map.get(s, "")} for s in leaf_ids],
            }
        return tasks
    except Exception:
        return {}

def _extract_json_object(text: str) -> Dict[str, Any]:
    clean = re.sub(r"^```(?:json)?\s*|\s*```$", "", str(text or "").strip(), flags=re.IGNORECASE)
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        start = clean.find("{")
        end = clean.rfind("}")
        if start >= 0 and end > start:
            return json.loads(clean[start:end + 1])
        raise

def _sections_to_markdown_content(data: Dict[str, Any]) -> str:
    parts = []
    for item in data.get("sections", []) or []:
        sid = str(item.get("node_id", "")).strip()
        title = str(item.get("title", "")).strip()
        md = str(item.get("markdown", "")).strip()
        if not sid or not md:
            continue
        level = min(len(sid.split(".")), 6)
        heading = f"{'#' * level} {sid} {title}".strip()
        parts.append(f"{heading}\n\n{md}")
    return "\n\n".join(parts).strip()

def _extract_pmids(markdown_text: str) -> List[str]:
    text = str(markdown_text or "")
    pmids: List[str] = []
    citation_pattern = re.compile(r'[（(]\s*PMID\s*[:：]([^）)]*)[）)]', re.IGNORECASE)
    spans = []
    for match in citation_pattern.finditer(text):
        spans.append(match.span())
        pmids.extend(re.findall(r'\b\d{5,}\b', match.group(1)))

    if spans:
        masked = list(text)
        for start, end in spans:
            masked[start:end] = [" "] * (end - start)
        text = "".join(masked)
    pmids.extend(re.findall(r'PMID\s*[:：]\s*(\d{5,})', text, flags=re.IGNORECASE))
    return pmids

def _normalize_pmid_citations(markdown_text: str) -> str:
    """Normalize citation punctuation before validation/export.

    Accepted model variants such as （pmid：123，456） or
    (PMID: 123, PMID: 456) are stored as (PMID: 123, 456).
    """
    citation_pattern = re.compile(r'[（(]\s*PMID\s*[:：]([^）)]*)[）)]', re.IGNORECASE)

    def repl(match):
        pmids = re.findall(r'\b\d{5,}\b', match.group(1))
        if not pmids:
            return match.group(0)
        return f"(PMID: {', '.join(pmids)})"

    return citation_pattern.sub(repl, str(markdown_text or ""))

def _validate_section_citations(data: Dict[str, Any], allowed_pmids_by_scope: Dict[str, set]) -> List[Dict[str, Any]]:
    problems = []
    for item in data.get("sections", []) or []:
        sid = str(item.get("node_id", "")).strip()
        allowed = allowed_pmids_by_scope.get(sid, set())
        used = set(_extract_pmids(item.get("markdown", "")))
        invalid = sorted(pmid for pmid in used if pmid not in allowed)
        if invalid:
            problems.append({"node_id": sid, "invalid_pmids": invalid, "allowed_pmids": sorted(allowed)})
    return problems

def _build_structured_writing_prompts(
    section_data: Dict[str, Any],
    global_topic: str,
    ref_blocks_by_scope: Dict[str, str],
    allowed_pmids_by_scope: Dict[str, set],
    lang: str = 'zh',
) -> Tuple[str, str]:
    sid = section_data["id"]
    style = "professional academic English" if lang == 'en' else "专业、客观、严谨的学术中文"
    topic_label = "Review topic" if lang == 'en' else "综述主题"
    title_label = "Current unit" if lang == 'en' else "当前写作单元"
    parent_label = "Parent section" if lang == 'en' else "父级节点"
    nodes_label = "Nodes in this unit" if lang == 'en' else "当前单元包含节点"
    sibling_label = "Sibling nodes for boundary reference only; do not expand them" if lang == 'en' else "同级节点，仅供边界参考，不要展开"
    scope_label = "Citation scopes" if lang == 'en' else "引用作用域"
    refs_label = "Reference materials" if lang == 'en' else "文献素材"
    req_label = "Writing requirements" if lang == 'en' else "写作要求"
    no_cite = "do not use PMID citations" if lang == 'en' else "不使用 PMID 引文"

    sections_schema = []
    unit_ids = []
    if section_data["id"] not in [n["serial"] for n in section_data.get("leaf_nodes", [])]:
        unit_ids.append(section_data["id"])
    unit_ids.extend([n["serial"] for n in section_data.get("leaf_nodes", [])])
    for node_id in unit_ids:
        sections_schema.append({
            "node_id": node_id,
            "title": section_data["title"] if node_id == section_data["id"] else next((n["title"] for n in section_data.get("leaf_nodes", []) if n["serial"] == node_id), ""),
            "markdown": "...",
        })

    system_prompt = (
        "You are a medical review writing assistant. Write only from the supplied outline and reference materials.\n"
        "Hard rules:\n"
        "1. Use only the supplied reference materials.\n"
        "2. Citations must use ASCII punctuation and the exact format (PMID: 12345678).\n"
        "   For multiple PMIDs, use exactly (PMID: 12345678, 23456789). Do not repeat the PMID label inside one citation.\n"
        "3. Each node_id may cite only PMIDs listed in its citation scope.\n"
        "4. Do not invent PMIDs.\n"
        "5. Do not output a reference list.\n"
        "6. Output valid JSON only, with no text outside JSON.\n"
    )

    parent_line = ""
    if section_data.get("parent_id"):
        parent_line = f"{section_data['parent_id']} {section_data.get('parent_title', '')}".strip()

    sibling_lines = [
        f"{n['serial']} {n['title']}".strip()
        for n in section_data.get("sibling_nodes", [])
    ]
    if not sibling_lines:
        sibling_lines = ["None" if lang == 'en' else "无"]

    scope_lines = []
    for node_id in unit_ids:
        allowed = sorted(allowed_pmids_by_scope.get(node_id, set()))
        scope_lines.append(f"{node_id}: {', '.join(allowed) if allowed else no_cite}")

    refs = "\n".join(ref_blocks_by_scope.get(n["serial"], "") for n in section_data.get("leaf_nodes", []) if ref_blocks_by_scope.get(n["serial"], "")).strip()

    req_lines = (
        [
            f"1. Write in {style}.",
            "2. Each sections[i].markdown must contain only the body text for its own node_id.",
            "3. Do not repeat the section number or title inside markdown.",
            "4. Do not expand sibling nodes listed above.",
            "5. Parent or unit-introduction nodes without citation scope should not contain PMID citations.",
            "6. Use only ASCII citation punctuation. For one citation use (PMID: 12345678); for multiple PMIDs use (PMID: 12345678, 23456789).",
            "7. Output valid JSON only.",
        ]
        if lang == 'en'
        else [
            f"1. 使用{style}。",
            "2. 每个 sections[i].markdown 只写对应 node_id 的正文。",
            "3. 不要在 markdown 中重复小节编号或标题。",
            "4. 不要展开上方列出的同级节点主题。",
            "5. 没有引用作用域的父节点或单元导语不要包含 PMID 引文。",
            "6. PMID 引文统一使用英文半角标点：单篇写作 (PMID: 12345678)，多篇写作 (PMID: 12345678, 23456789)。",
            "7. 只输出合法 JSON。",
        ]
    )

    user_prompt = f"""{topic_label}:
{global_topic}

{parent_label}:
{parent_line or ("None" if lang == 'en' else "无")}

{title_label}:
{sid} {section_data.get('title', '')}

{nodes_label}:
{chr(10).join(section_data.get('skeleton', []))}

{sibling_label}:
{chr(10).join(sibling_lines)}

{scope_label}:
{chr(10).join(scope_lines)}

JSON output shape:
{json.dumps({"unit_id": sid, "sections": sections_schema}, ensure_ascii=False, indent=2)}

{refs_label}:
{refs}

{req_label}:
{chr(10).join(req_lines)}"""
    return system_prompt, user_prompt

def _build_repair_prompts(
    unit_id: str,
    item: Dict[str, Any],
    invalid_pmids: List[str],
    ref_block: str,
    allowed_pmids: set,
    lang: str = 'zh',
) -> Tuple[str, str]:
    node_id = str(item.get("node_id", "")).strip()
    title = str(item.get("title", "")).strip()
    no_cite = "No PMID citations are allowed for this node." if lang == 'en' else "该节点不允许使用 PMID 引文。"
    allowed_line = ", ".join(sorted(allowed_pmids)) if allowed_pmids else no_cite
    style = "professional academic English" if lang == 'en' else "专业、客观、严谨的学术中文"
    system_prompt = (
        "You are a medical review writing assistant repairing one subsection.\n"
        "Output valid JSON only. Do not output text outside JSON. Do not invent PMIDs. "
        "Use ASCII citation punctuation only: (PMID: 12345678) or (PMID: 12345678, 23456789)."
    )
    user_prompt = f"""Rewrite only this subsection to remove citation-scope drift.

Unit ID:
{unit_id}

Node:
{node_id} {title}

Invalid PMIDs that must not appear:
{', '.join(invalid_pmids)}

Allowed PMIDs for this node:
{allowed_line}

Original markdown:
{item.get('markdown', '')}

Reference materials for this node:
{ref_block or 'None'}

Requirements:
1. Write in {style}.
2. Keep the same node_id and title.
3. The markdown must cite only the allowed PMIDs listed above.
4. If no PMID is allowed or no suitable material supports a claim, write without PMID citation.
5. Use only ASCII citation punctuation. For one citation use (PMID: 12345678); for multiple PMIDs use (PMID: 12345678, 23456789).
6. Do not repeat the section number or title inside markdown.

JSON output shape:
{{
  "node_id": "{node_id}",
  "title": {json.dumps(title, ensure_ascii=False)},
  "markdown": "..."
}}"""
    return system_prompt, user_prompt

async def generate_section_content(collection_name: str, section_data: Dict, global_topic: str, rag_k: int, llm_config: Dict, embedding_config: Dict, log_callback, lang='zh') -> Dict:
    sid, stitle = section_data['id'], section_data['title']
    log_callback(f"Processing writing unit {sid}: {stitle}...")
    sem = asyncio.Semaphore(5)
    async def get_ref(node):
        async with sem: return node, await search_text_from_chromadb(collection_name, f"{stitle} - {node['title']}", rag_k, embedding_config)
    results = await asyncio.gather(*[get_ref(n) for n in section_data['leaf_nodes']])
    
    ref_blocks_by_scope: Dict[str, str] = {}
    allowed_pmids_by_scope: Dict[str, set] = {section_data["id"]: set()}
    reference_materials_by_scope: Dict[str, Dict[str, List[str]]] = {}
    label_ref = "Ref for" if lang == 'en' else "针对标题"
    label_mat = "Material" if lang == 'en' else "素材"
    
    for node, res_list in results:
        node_id = node["serial"]
        allowed_pmids_by_scope.setdefault(node_id, set())
        reference_materials_by_scope.setdefault(node_id, {})
        if not res_list:
            continue
        lines = [f"\n[{label_ref} {node_id} {node['title']}]:"] 
        for i, r in enumerate(res_list):
            pmid = str(r.get('pmid', '')).strip()
            content = str(r.get('content', ''))[:3000]
            if pmid:
                allowed_pmids_by_scope[node_id].add(pmid)
                reference_materials_by_scope[node_id].setdefault(pmid, [])
                if content and content not in reference_materials_by_scope[node_id][pmid]:
                    reference_materials_by_scope[node_id][pmid].append(content)
            lines.append(f"- {label_mat} {i+1} (PMID: {pmid}): {content}")
        ref_blocks_by_scope[node_id] = "\n".join(lines)

    system_prompt, prompt = _build_structured_writing_prompts(
        section_data,
        global_topic,
        ref_blocks_by_scope,
        allowed_pmids_by_scope,
        lang=lang,
    )

    output_dir = get_review_parts_dir(collection_name, lang)
    os.makedirs(output_dir, exist_ok=True)
    
    try:
        prompt_file = os.path.join(output_dir, f"prompt_section_{sid}.txt")
        with open(prompt_file, "w", encoding="utf-8") as f:
            f.write("[SYSTEM]\n")
            f.write(system_prompt)
            f.write("\n\n[USER]\n")
            f.write(prompt)
    except: pass

    log_callback(f"  - Requesting LLM for {sid}...")
    res = await call_llm_openai(prompt, llm_config, log_callback, system_prompt=system_prompt)
    if not res:
        log_callback(f"× Writing unit {sid} failed: LLM returned empty response.")
        return {"status": "failed"}

    try:
        data = _extract_json_object(res)
        if "sections" not in data:
            legacy_md = str(data.get("markdown_content", "")).strip()
            if legacy_md:
                data["sections"] = [{"node_id": sid, "title": stitle, "markdown": legacy_md}]
        data["unit_id"] = str(data.get("unit_id") or sid)
        for item in data.get("sections", []) or []:
            item["markdown"] = _normalize_pmid_citations(item.get("markdown", ""))

        problems = _validate_section_citations(data, allowed_pmids_by_scope)
        if problems:
            log_callback(f"  - Citation drift detected in {sid}: {len(problems)} subsection(s). Repairing...")
            section_by_id = {str(item.get("node_id", "")).strip(): item for item in data.get("sections", [])}
            for problem in problems:
                node_id = problem["node_id"]
                item = section_by_id.get(node_id)
                if not item:
                    continue
                repair_system, repair_prompt = _build_repair_prompts(
                    sid,
                    item,
                    problem["invalid_pmids"],
                    ref_blocks_by_scope.get(node_id, ""),
                    allowed_pmids_by_scope.get(node_id, set()),
                    lang=lang,
                )
                repair_res = await call_llm_openai(repair_prompt, llm_config, log_callback, system_prompt=repair_system)
                if not repair_res:
                    continue
                try:
                    repaired_item = _extract_json_object(repair_res)
                    if str(repaired_item.get("node_id", "")).strip() == node_id:
                        item["markdown"] = _normalize_pmid_citations(str(repaired_item.get("markdown", "")).strip())
                except Exception as repair_error:
                    log_callback(f"  - Repair parse failed for {node_id}: {repair_error}")

        final_problems = _validate_section_citations(data, allowed_pmids_by_scope)
        data["citation_validation"] = {
            "status": "passed" if not final_problems else "failed",
            "problems": final_problems,
        }
        data["citation_scopes"] = {k: sorted(v) for k, v in allowed_pmids_by_scope.items()}
        data["reference_materials"] = reference_materials_by_scope
        data["markdown_content"] = _sections_to_markdown_content(data) or str(data.get("markdown_content", "")).strip()
        os.makedirs(output_dir, exist_ok=True)
        output_file = os.path.join(output_dir, f"section_{sid}.json")
        with open(output_file, "w", encoding="utf-8") as f: json.dump(data, f, indent=2, ensure_ascii=False)
        if final_problems:
            log_callback(f"× Writing unit {sid} saved, but citation validation still failed.")
            return {"status": "validation_failed", "file": output_file, "problems": final_problems}
        log_callback(f"√ Writing unit {sid} done and citations validated.")
        return {"status": "success", "file": output_file}
    except Exception as e:
        log_callback(f"× Writing unit {sid} failed: {e}")
        return {"status": "failed", "error": str(e)}

# ==========================================
#               6. Export Helpers
# ==========================================

def _merge_reference_entry(section_ref_db: Dict[str, Dict[str, List[str]]], scope: str, pmid: str, content: str):
    clean_scope = str(scope or "").strip()
    clean_pmid = str(pmid or "").strip()
    clean_content = str(content or "").strip()
    if not clean_scope or not clean_pmid or not clean_content:
        return
    section_ref_db.setdefault(clean_scope, {})
    section_ref_db[clean_scope].setdefault(clean_pmid, [])
    if clean_content not in section_ref_db[clean_scope][clean_pmid]:
        section_ref_db[clean_scope][clean_pmid].append(clean_content)

def _extract_references_from_section_jsons(parts_dir: str) -> Dict[str, Dict[str, List[str]]]:
    section_ref_db = {}
    if not os.path.exists(parts_dir):
        return section_ref_db

    for jf in glob.glob(os.path.join(parts_dir, "section_*.json")):
        try:
            with open(jf, 'r', encoding='utf-8') as f:
                data = json.load(f)
            ref_materials = data.get("reference_materials", {})
            if not isinstance(ref_materials, dict):
                continue
            for scope, pmids in ref_materials.items():
                if not isinstance(pmids, dict):
                    continue
                for pmid, entries in pmids.items():
                    if isinstance(entries, list):
                        for content in entries:
                            _merge_reference_entry(section_ref_db, scope, pmid, content)
                    else:
                        _merge_reference_entry(section_ref_db, scope, pmid, entries)
        except Exception:
            pass
    return section_ref_db

def _is_prompt_reference_boundary(line: str, header_pattern, ref_pattern) -> bool:
    stripped = str(line or "").strip()
    if not stripped:
        return False
    boundary_prefixes = (
        "### ",
        "[SYSTEM]",
        "[USER]",
        "写作要求:",
        "Writing requirements:",
        "Requirements:",
        "JSON output shape:",
        "引用作用域:",
        "Citation scopes:",
    )
    return (
        header_pattern.search(stripped)
        or ref_pattern.match(stripped)
        or any(stripped.startswith(prefix) for prefix in boundary_prefixes)
    )

def _extract_references_from_prompt_files(parts_dir: str) -> Dict[str, Dict[str, List[str]]]:
    section_ref_db = {}
    if not os.path.exists(parts_dir):
        return section_ref_db
    prompt_files = glob.glob(os.path.join(parts_dir, "prompt_section_*.txt"))
    section_id_pattern = r'[A-Za-z0-9]+(?:\.[A-Za-z0-9]+)*'
    header_pattern = re.compile(r'\[(?:针对标题|Ref for)\s*(' + section_id_pattern + r')', re.IGNORECASE)
    ref_pattern = re.compile(r'-\s*(?:素材|Material)\s*\d+\s*\(PMID:\s*(\d+)\):\s*(.*)', re.IGNORECASE)

    for p_file in prompt_files:
        try:
            filename = os.path.basename(p_file)
            file_main_id = filename.replace("prompt_section_", "").replace(".txt", "")
            current_scope_id = file_main_id
            if current_scope_id not in section_ref_db: section_ref_db[current_scope_id] = {}

            with open(p_file, 'r', encoding='utf-8') as f: lines = [line.rstrip('\n') for line in f]
            idx = 0
            while idx < len(lines):
                raw_line = lines[idx]
                line = raw_line.strip()
                if not line:
                    idx += 1
                    continue
                h_match = header_pattern.search(line)
                if h_match:
                    current_scope_id = h_match.group(1)
                    if current_scope_id not in section_ref_db: section_ref_db[current_scope_id] = {}
                    idx += 1
                    continue
                m_match = ref_pattern.match(line)
                if m_match:
                    pmid = m_match.group(1)
                    content_lines = [m_match.group(2).rstrip()]
                    idx += 1
                    while idx < len(lines):
                        next_raw = lines[idx]
                        if _is_prompt_reference_boundary(next_raw, header_pattern, ref_pattern):
                            break
                        content_lines.append(next_raw.rstrip())
                        idx += 1
                    content = "\n".join(content_lines).strip()
                    _merge_reference_entry(section_ref_db, current_scope_id, pmid, content)
                    continue
                idx += 1
        except: pass
    return section_ref_db

def extract_references_from_prompts(parts_dir: str) -> Dict[str, Dict[str, List[str]]]:
    """Build the HTML reference database.

    New writing outputs store reference materials in section_*.json. Prompt parsing remains
    a compatibility fallback for older review_parts folders.
    """
    section_ref_db = _extract_references_from_section_jsons(parts_dir)
    prompt_ref_db = _extract_references_from_prompt_files(parts_dir)
    for scope, pmids in prompt_ref_db.items():
        for pmid, entries in pmids.items():
            for content in entries:
                _merge_reference_entry(section_ref_db, scope, pmid, content)
    return section_ref_db

def strip_leading_section_heading(markdown_text: str, sid: str) -> str:
    """Remove the section's own leading markdown heading while keeping the body intact."""
    text = str(markdown_text or "")
    lines = text.splitlines()
    first_nonempty = next((idx for idx, line in enumerate(lines) if line.strip()), None)
    if first_nonempty is None:
        return ""

    heading_pattern = re.compile(r'^\s*#{1,6}\s+' + re.escape(str(sid).strip()) + r'(?:\s+.*)?\s*$')
    if heading_pattern.match(lines[first_nonempty]):
        del lines[first_nonempty]
        while first_nonempty < len(lines) and not lines[first_nonempty].strip():
            del lines[first_nonempty]

    return "\n".join(lines).strip()


def infer_review_section_ref_scopes(task_leaf_ids: List[str], sid: str) -> List[str]:
    """Use the leaf nodes that fed the prompt as the citation lookup scope for a rendered section."""
    clean_sid = str(sid or "").strip()
    clean_leaf_ids = [str(leaf).strip() for leaf in task_leaf_ids if str(leaf).strip()]

    descendant_leaf_ids = [leaf for leaf in clean_leaf_ids if leaf == clean_sid or leaf.startswith(clean_sid + ".")]
    if descendant_leaf_ids:
        return descendant_leaf_ids
    if clean_leaf_ids:
        return clean_leaf_ids
    return [clean_sid] if clean_sid else []


def load_normalized_review_sections(parts_dir: str, csv_path: str, log_callback=None) -> Dict[str, Dict[str, Any]]:
    """Normalize review_parts into per-section bodies plus citation scopes shared by HTML and DOCX exporters."""
    normalized_sections: Dict[str, Dict[str, Any]] = {}
    if not os.path.exists(parts_dir) or not os.path.exists(csv_path):
        return normalized_sections

    tasks = parse_framework_to_sections(csv_path)
    current_task_files = {
        os.path.basename(path).replace("section_", "").replace(".json", "")
        for path in glob.glob(os.path.join(parts_dir, "section_*.json"))
        if os.path.basename(path).replace("section_", "").replace(".json", "") in tasks
    }
    header_pattern = re.compile(r'(^|\n)(#+)\s+([\d\.]+)(\s+|$)', re.MULTILINE)

    def append_section(render_sid: str, markdown_text: str, ref_scopes: List[str]):
        clean_sid = str(render_sid or "").strip()
        clean_text = str(markdown_text or "").strip()
        if not clean_sid or not clean_text:
            return

        payload = normalized_sections.setdefault(clean_sid, {"markdown_parts": [], "reference_scopes": []})
        payload["markdown_parts"].append(clean_text)
        for scope in ref_scopes:
            clean_scope = str(scope or "").strip()
            if clean_scope and clean_scope not in payload["reference_scopes"]:
                payload["reference_scopes"].append(clean_scope)

    for jf in sorted(glob.glob(os.path.join(parts_dir, "section_*.json"))):
        try:
            main_id = os.path.basename(jf).replace("section_", "").replace(".json", "")
            with open(jf, 'r', encoding='utf-8') as f:
                section_json = json.load(f)

            structured_sections = section_json.get("sections")
            if isinstance(structured_sections, list) and structured_sections:
                for item in structured_sections:
                    sid = str(item.get("node_id", "")).strip()
                    if main_id != sid and sid in current_task_files:
                        continue
                    append_section(
                        sid,
                        str(item.get("markdown", "")).strip(),
                        [sid],
                    )
                continue

            raw = section_json.get('markdown_content', '')
            raw = str(raw or "")

            task = tasks.get(main_id, {})
            task_leaf_ids = [
                str(node.get("serial", "")).strip()
                for node in task.get("leaf_nodes", [])
                if str(node.get("serial", "")).strip()
            ]

            matches = list(header_pattern.finditer(raw))
            if not matches:
                target_sid = main_id
                if len(task_leaf_ids) == 1:
                    target_sid = task_leaf_ids[0]
                    if log_callback and target_sid != main_id:
                        log_callback(f"Section {main_id}: reassigned untitled content to sole leaf {target_sid}")
                elif len(task_leaf_ids) > 1 and log_callback:
                    log_callback(f"Section {main_id}: untitled multi-leaf content kept at parent {main_id}")

                append_section(
                    target_sid,
                    strip_leading_section_heading(raw, target_sid),
                    infer_review_section_ref_scopes(task_leaf_ids, target_sid),
                )
                continue

            prefix = raw[:matches[0].start()].strip()
            if prefix and main_id in tasks:
                append_section(main_id, prefix, infer_review_section_ref_scopes(task_leaf_ids, main_id))

            for i, match in enumerate(matches):
                sid = match.group(3).strip()
                if main_id != sid and sid in current_task_files:
                    continue
                end = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
                chunk = raw[match.start():end].strip()
                append_section(
                    sid,
                    strip_leading_section_heading(chunk, sid),
                    infer_review_section_ref_scopes(task_leaf_ids, sid),
                )
        except Exception:
            pass

    finalized_sections = {}
    for sid, payload in normalized_sections.items():
        markdown_text = "\n\n".join(part for part in payload["markdown_parts"] if part.strip()).strip()
        if not markdown_text:
            continue
        finalized_sections[sid] = {
            "markdown": markdown_text,
            "reference_scopes": payload["reference_scopes"],
        }
    return finalized_sections


def get_all_references_metadata(collection_name: str) -> Dict[str, Dict]:
    c_dir = get_collection_path(collection_name)
    xmls = glob.glob(os.path.join(c_dir, "*.xml"))
    ref_map = {}
    if xmls:
        try:
            root = ET.parse(xmls[0]).getroot()
            for art in root.findall('.//PubmedArticle'):
                pmid_elem = art.find('./MedlineCitation/PMID')
                title_elem = art.find('./MedlineCitation/Article/ArticleTitle')
                if pmid_elem is None or title_elem is None:
                    continue
                pmid = pmid_elem.text
                title = "".join(title_elem.itertext())
                doi = next((x.text for x in art.findall('./PubmedData/ArticleIdList/ArticleId') if x.get('IdType') == 'doi'), None)
                auth = "Unknown"
                al = art.findall('./MedlineCitation/Article/AuthorList/Author')
                if al: auth = f"{al[0].find('LastName').text} {al[0].find('Initials').text} et al."
                ref_map[pmid] = {"title": title, "doi": doi, "author": auth}
        except: pass
    return ref_map

def load_and_split_all_sections(parts_dir: str) -> Dict[str, str]:
    content_map = {}
    if not os.path.exists(parts_dir): return content_map
    header_pattern = re.compile(r'(^|\n)(#+)\s+([\d\.]+)(\s+|$)', re.MULTILINE)
    for jf in glob.glob(os.path.join(parts_dir, "section_*.json")):
        try:
            main_id = os.path.basename(jf).replace("section_", "").replace(".json", "")
            with open(jf, 'r', encoding='utf-8') as f:
                section_json = json.load(f)
            structured_sections = section_json.get("sections")
            if isinstance(structured_sections, list) and structured_sections:
                for item in structured_sections:
                    sid = str(item.get("node_id", "")).strip()
                    if sid:
                        content_map[sid] = str(item.get("markdown", "")).strip()
                continue
            raw = section_json.get('markdown_content', '')
            matches = list(header_pattern.finditer(raw))
            if not matches: content_map[main_id] = raw; continue
            if matches[0].start() > 0: content_map[main_id] = content_map.get(main_id, "") + raw[:matches[0].start()]
            for i, m in enumerate(matches):
                sid = m.group(3)
                end = matches[i+1].start() if i+1 < len(matches) else len(raw)
                content_map[sid] = raw[m.start():end].strip()
        except: pass
    return content_map

async def generate_final_word_doc(collection_name: str, topic: str, log_callback, lang='zh') -> str:
    c_dir = get_collection_path(collection_name)
    parts = get_review_parts_dir(collection_name, lang, allow_legacy=True)
    csv = get_framework_csv_path(collection_name, lang, allow_legacy=True)
    out = os.path.join(parts, f"{collection_name}_Review.docx")
    if not os.path.exists(parts) or not os.path.exists(csv): return None
    txt_time, txt_ref, txt_no_meta = ("Generated Time", "References", "Metadata not found") if lang == 'en' else ("生成时间", "参考文献", "未找到详细元数据")
    doc = Document()
    try:
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    except: pass
    t_para = doc.add_heading(topic, 0)
    t_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"{txt_time}: {time.strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()
    df = pd.read_csv(csv, dtype=str)
    df.columns = ['Serial', 'Title'] + list(df.columns[2:])
    rows = sorted(df.to_dict('records'), key=lambda x: [int(y) if y.isdigit() else 0 for y in str(x['Serial']).split('.')])
    section_map, pmids = load_normalized_review_sections(parts, csv, log_callback), set()
    for r in rows:
        s = str(r['Serial'])
        doc.add_heading(f"{s} {r['Title']}", level=min(len(s.split('.')), 9))
        section_payload = section_map.get(s)
        if section_payload:
            md = section_payload["markdown"]
            for l in md.split('\n'):
                if not l.strip() or l.startswith('#'): continue
                para = doc.add_paragraph()
                for pt in re.split(r'(\*\*.*?\*\*)', l):
                    run = para.add_run(pt.replace('**', ''))
                    if pt.startswith('**'): run.bold = True
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            pmids.update(_extract_pmids(md))
    if pmids:
        doc.add_page_break()
        doc.add_heading(txt_ref, level=1)
        ref_map = get_all_references_metadata(collection_name)
        for pmid in sorted(list(pmids)):
            info = ref_map.get(pmid)
            p = doc.add_paragraph()
            p.add_run(f"[{pmid}] ").bold = True
            if info:
                p.add_run(f"{info['author']}. ")
                p.add_run(f"{info['title']}").italic = True
                if info['doi']: p.add_run(f" (DOI: {info['doi']})")
            else: p.add_run(f"({txt_no_meta})")
    doc.save(out); return out

async def generate_interactive_html_review(collection_name: str, topic: str, log_callback, lang='zh') -> str:
    c_dir = get_collection_path(collection_name)
    parts = get_review_parts_dir(collection_name, lang, allow_legacy=True)
    csv = get_framework_csv_path(collection_name, lang, allow_legacy=True)
    out = os.path.join(parts, f"{collection_name}_Review.html")
    if not os.path.exists(parts) or not os.path.exists(csv): return None
    txt_toc, txt_time, txt_ref_title, txt_tip, txt_source, txt_inherit, txt_root, txt_goto, txt_not_found, txt_ctx = ("Table of Contents", "Generated Time", "References", "Click blue PMIDs in text to view details", "Source Section", "Inherited from", "Root Section", "Go to Pubmed", "Material not found", "Context") if lang == 'en' else ("目录", "生成时间", "引用详情", "点击正文中蓝色的 PMID<br>查看该章节参考的素材详情", "来源章节", "继承自章节", "根章节", "前往 Pubmed", "未找到对应素材", "当前上下文")
    ref_db = extract_references_from_prompts(parts)
    df = pd.read_csv(csv, dtype=str)
    df.columns = ['Serial', 'Title'] + list(df.columns[2:])
    rows = sorted(df.to_dict('records'), key=lambda x: [int(y) if y.isdigit() else 0 for y in str(x['Serial']).split('.')])
    section_map, toc_html, body_html = load_normalized_review_sections(parts, csv, log_callback), "", ""
    import markdown
    pmid_re = re.compile(r'([（\(]\s*PMID\s*[:：]\s*)([^）\)]+)([）\)])', re.IGNORECASE)
    for r in rows:
        s, t = str(r['Serial']), str(r['Title'])
        toc_html += f'<div class="toc-item" style="padding-left:{(len(s.split('.'))-1)*20}px"><a href="#s-{s}">{s} {t}</a></div>'
        body_html += f'<div id="s-{s}" class="sec"><h{min(len(s.split('.')), 6)}>{s} {t}</h{min(len(s.split('.')), 6)}>'
        section_payload = section_map.get(s)
        if section_payload:
            scope_json = json.dumps(section_payload.get("reference_scopes") or [s], ensure_ascii=False)
            html = markdown.markdown(section_payload["markdown"])
            def repl(m):
                pre, inner, suf = m.groups()
                links = [f"<span class=\"lnk\" onclick='show(\"{i}\", {scope_json})'>{i}</span>" for i in re.findall(r'\d+', inner)]
                return pre + ", ".join(links) + suf
            body_html += f'<div class="content">{pmid_re.sub(repl, html)}</div>'
        body_html += '</div><hr/>'
    js_data = json.dumps(ref_db, ensure_ascii=False)
    script_html = f"""
<script>
const db = {js_data};
const TXT_SOURCE = {json.dumps(txt_source, ensure_ascii=False)};
const TXT_INHERIT = {json.dumps(txt_inherit, ensure_ascii=False)};
const TXT_ROOT = {json.dumps(txt_root, ensure_ascii=False)};
const TXT_GOTO = {json.dumps(txt_goto, ensure_ascii=False)};
const TXT_NOT_FOUND = {json.dumps(txt_not_found, ensure_ascii=False)};
const TXT_CTX = {json.dumps(txt_ctx, ensure_ascii=False)};

function lineage(scope) {{
  const items = [];
  let parts = String(scope || '').split('.');
  while (parts.length > 0) {{
    items.push(parts.join('.'));
    parts.pop();
  }}
  return items;
}}

function relation(origin, matched) {{
  const originText = String(origin || '');
  if (matched === originText) return TXT_SOURCE + ': ' + matched;
  if (matched === originText.split('.')[0]) return TXT_ROOT + ': ' + matched;
  return TXT_INHERIT + ': ' + matched;
}}

function show(id, scopes) {{
  const area = document.getElementById('area');
  const scopeList = Array.isArray(scopes) ? scopes : [scopes];
  const cards = [];
  const seen = new Set();

  for (const origin of scopeList) {{
    if (!origin) continue;
    for (const matched of lineage(origin)) {{
      const entries = db[matched] && db[matched][id];
      if (!entries) continue;
      const msg = relation(origin, matched);
      for (const text of entries) {{
        const key = matched + '|' + text;
        if (seen.has(key)) continue;
        seen.add(key);
        cards.push(
          '<div class="card"><div class="ct">PMID: ' + id + ' (' + msg + ')</div><div class="cb">' + text + '</div>' +
          '<div style="margin-top:10px;font-size:12px"><a href="https://pubmed.ncbi.nlm.nih.gov/' + id + '/" target="_blank">🔗 ' + TXT_GOTO + '</a></div></div>'
        );
      }}
    }}
  }}

  if (cards.length) {{
    area.innerHTML = cards.join('');
    return;
  }}

  const context = scopeList.filter(Boolean).join(', ');
  area.innerHTML =
    '<div class="card" style="border-left-color:#999"><div class="ct" style="color:#666">PMID: ' + id + '</div>' +
    '<div class="cb">' + TXT_NOT_FOUND + '<br>' + TXT_CTX + ': ' + (context || 'N/A') + '</div>' +
    '<div style="margin-top:10px;font-size:12px"><a href="https://pubmed.ncbi.nlm.nih.gov/' + id + '/" target="_blank">🔗 ' + TXT_GOTO + '</a></div></div>';
}}
</script>"""
    style_html = """
<style>
:root {
  --bg: #efe7da;
  --panel: #f7f3eb;
  --panel-strong: #f1ebdf;
  --paper: #fffdf9;
  --paper-soft: #fffaf2;
  --ink: #1f2933;
  --ink-soft: #5b6673;
  --line: #ded6c9;
  --line-strong: #cfc3b0;
  --accent: #0f6c5c;
  --accent-soft: #dff1eb;
  --accent-warm: #b86f3c;
  --shadow: 0 18px 42px rgba(83, 66, 44, 0.10);
}

* { box-sizing: border-box; }

html, body {
  margin: 0;
  padding: 0;
  height: 100%;
}

body {
  display: flex;
  overflow: hidden;
  color: var(--ink);
  background:
    radial-gradient(circle at top left, rgba(184, 111, 60, 0.10), transparent 26%),
    radial-gradient(circle at top right, rgba(15, 108, 92, 0.09), transparent 24%),
    linear-gradient(180deg, #f3ede3 0%, #ede5d8 100%);
  font-family: "Avenir Next", "PingFang SC", "Hiragino Sans GB", "Segoe UI", sans-serif;
}

#toc,
#ref {
  flex-shrink: 0;
  overflow-y: auto;
  padding: 22px 18px 28px;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.78) 0%, rgba(248, 244, 236, 0.92) 100%);
  backdrop-filter: blur(8px);
}

#toc {
  width: 250px;
  border-right: 1px solid var(--line);
}

#ref {
  width: 320px;
  border-left: 1px solid var(--line);
  display: flex;
  flex-direction: column;
}

#toc h3,
#ref h3 {
  margin: 0 0 16px;
  padding: 0 0 12px;
  position: sticky;
  top: 0;
  z-index: 2;
  background: inherit;
  color: var(--accent);
  font-size: 15px;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
  border-bottom: 1px solid var(--line);
}

.toc-item {
  margin-bottom: 6px;
  font-size: 13px;
  line-height: 1.45;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}

.toc-item a {
  display: block;
  padding: 7px 10px;
  border-radius: 10px;
  color: var(--ink);
  text-decoration: none;
  transition: background-color 120ms ease, color 120ms ease, transform 120ms ease;
}

.toc-item a:hover {
  color: var(--accent);
  background: rgba(15, 108, 92, 0.09);
  transform: translateX(2px);
}

#main {
  flex: 1;
  overflow-y: auto;
  scroll-behavior: smooth;
  padding: 44px 48px 56px;
  background:
    linear-gradient(180deg, rgba(255, 253, 249, 0.92) 0%, rgba(255, 250, 242, 0.96) 100%);
}

#main > h1 {
  margin: 0 0 10px;
  color: #14202b;
  font-size: clamp(34px, 3vw, 42px);
  line-height: 1.14;
  letter-spacing: -0.03em;
  font-weight: 750;
  max-width: 980px;
}

#main > p {
  margin: 0 0 34px;
  color: var(--ink-soft);
  font-size: 14px;
  letter-spacing: 0.02em;
}

.sec {
  max-width: 980px;
  margin: 0 auto;
}

.sec > h1,
.sec > h2,
.sec > h3,
.sec > h4,
.sec > h5,
.sec > h6 {
  margin: 0;
  color: #16222d;
  line-height: 1.28;
  font-weight: 720;
}

.sec > h1 {
  font-size: 29px;
  letter-spacing: -0.025em;
}

.sec > h2 {
  font-size: 24px;
  letter-spacing: -0.02em;
}

.sec > h3 {
  font-size: 20px;
}

.sec > h4,
.sec > h5,
.sec > h6 {
  font-size: 17px;
}

.content {
  margin-top: 16px;
  padding: 22px 26px;
  background: linear-gradient(180deg, var(--paper) 0%, var(--paper-soft) 100%);
  border: 1px solid rgba(207, 195, 176, 0.70);
  border-radius: 18px;
  box-shadow: var(--shadow);
}

.content p,
.content li,
.cb {
  font-family: "Iowan Old Style", "Palatino Linotype", "Book Antiqua", Georgia, serif;
}

.content p,
.content li {
  margin: 0 0 14px;
  color: #25313d;
  font-size: 17px;
  line-height: 1.9;
}

.content p:last-child,
.content li:last-child {
  margin-bottom: 0;
}

.content strong {
  color: #11202c;
}

.content ul,
.content ol {
  margin: 0 0 16px 1.4em;
  padding: 0;
}

.lnk {
  display: inline-block;
  margin: 0 2px;
  padding: 1px 7px;
  border-radius: 999px;
  color: var(--accent);
  cursor: pointer;
  font-weight: 700;
  font-size: 0.95em;
  background: var(--accent-soft);
  border: 1px solid rgba(15, 108, 92, 0.16);
  transition: transform 120ms ease, background-color 120ms ease, border-color 120ms ease;
}

.lnk:hover {
  transform: translateY(-1px);
  background: #cde8de;
  border-color: rgba(15, 108, 92, 0.28);
  text-decoration: none;
}

hr {
  max-width: 980px;
  margin: 24px auto 30px;
  border: 0;
  border-top: 1px solid rgba(184, 169, 145, 0.70);
}

.card {
  background: linear-gradient(180deg, #fffefb 0%, #f9f4eb 100%);
  border: 1px solid rgba(205, 193, 172, 0.82);
  border-left: 4px solid var(--accent);
  padding: 16px 16px 14px;
  border-radius: 14px;
  margin-bottom: 16px;
  box-shadow: 0 14px 28px rgba(73, 58, 39, 0.10);
}

.ct {
  margin-bottom: 10px;
  color: var(--accent);
  font-size: 13px;
  font-weight: 800;
  letter-spacing: 0.03em;
}

.cb {
  color: #394453;
  font-size: 15px;
  line-height: 1.75;
  white-space: pre-wrap;
  overflow-wrap: anywhere;
}

.card a {
  color: var(--accent-warm);
  text-decoration: none;
  font-weight: 600;
}

.card a:hover {
  text-decoration: underline;
}

.tip {
  margin-top: 56px;
  padding: 20px 14px;
  border: 1px dashed rgba(184, 169, 145, 0.9);
  border-radius: 14px;
  color: #7a6e60;
  text-align: center;
  line-height: 1.7;
  background: rgba(255, 250, 242, 0.72);
}

@media (max-width: 1200px) {
  #main { padding: 36px 32px 48px; }
  .content { padding: 20px 22px; }
}
</style>"""
    full_html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>{topic}</title>{style_html}</head><body><div id="toc"><h3>{txt_toc}</h3>{toc_html}</div><div id="main"><h1>{topic}</h1><p style="text-align:center;color:#666">{txt_time}: {time.strftime('%Y-%m-%d %H:%M')}</p>{body_html}</div><div id="ref"><h3>{txt_ref_title}</h3><div id="area"><div class="tip">{txt_tip}</div></div></div>{script_html}</body></html>"""
    with open(out, "w", encoding="utf-8") as f: f.write(full_html)
    return out
